from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
IMG = ROOT / "images"


def set_run(run, size: float = 14, bold: bool = False, italic: bool = False, mono: bool = False) -> None:
    name = "Courier New" if mono else "Times New Roman"
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def fmt(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=True, before=0, after=0, line=1.0) -> None:
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(1.25) if first else Cm(0)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def p(doc: Document, text: str = "", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=True, bold=False, italic=False, size=14, before=0, after=0) -> None:
    par = doc.add_paragraph()
    fmt(par, align=align, first=first, before=before, after=after)
    run = par.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)


def h(doc: Document, text: str) -> None:
    par = doc.add_paragraph()
    par.style = doc.styles["Heading 1"]
    fmt(par, align=WD_ALIGN_PARAGRAPH.LEFT, first=True, before=10, after=6)
    run = par.add_run(text)
    set_run(run, bold=True)


def sh(doc: Document, text: str) -> None:
    par = doc.add_paragraph()
    par.style = doc.styles["Heading 2"]
    fmt(par, align=WD_ALIGN_PARAGRAPH.LEFT, first=True, before=8, after=4)
    run = par.add_run(text)
    set_run(run, bold=True)


def caption(doc: Document, text: str) -> None:
    par = doc.add_paragraph()
    fmt(par, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, before=2, after=7)
    run = par.add_run(text)
    set_run(run)


def cell(cell, text: str, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=14) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    par = cell.paragraphs[0]
    fmt(par, align=align, first=False, before=0, after=0)
    run = par.add_run(text)
    set_run(run, size=size, bold=bold)


def table(doc: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl.autofit = False
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cl = tbl.cell(r, c)
            if widths:
                cl.width = Cm(widths[c])
            cell(cl, text, bold=(r == 0), align=WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    p(doc, "", first=False, after=0)


def pic(doc: Document, filename: str, text: str, width: float = 15.6) -> None:
    par = doc.add_paragraph()
    fmt(par, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, before=4, after=2)
    par.add_run().add_picture(str(IMG / filename), width=Cm(width))
    caption(doc, text)


def code(doc: Document, title: str, path: Path) -> None:
    sh(doc, title)
    for line in path.read_text(encoding="utf-8").splitlines():
        par = doc.add_paragraph()
        fmt(par, align=WD_ALIGN_PARAGRAPH.LEFT, first=False, before=0, after=0, line=1.0)
        run = par.add_run(line)
        set_run(run, size=8.5, mono=True)


def add_page_number(section) -> None:
    par = section.footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    set_run(run, size=12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    add_page_number(section)
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.color.rgb = RGBColor(0, 0, 0)
        if name.startswith("Heading"):
            style.font.bold = True


def title_page(doc: Document) -> None:
    lines = [
        "Министерство образования Республики Беларусь",
        "",
        "Учреждение образования",
        "БЕЛОРУССКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ",
        "ИНФОРМАТИКИ И РАДИОЭЛЕКТРОНИКИ",
        "",
        "Факультет компьютерных систем и сетей",
        "",
        "Кафедра электронных вычислительных машин",
        "",
        "Дисциплина: Автоматизированное проектирование цифровых устройств",
    ]
    for line in lines:
        p(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    for _ in range(5):
        p(doc, "", first=False)
    p(doc, "ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ №2", align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True)
    p(doc, "на тему", align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    p(doc, "«Использование модулей памяти lpm_rom и lpm_ram_io»", align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True)
    p(doc, "Вариант 4", align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    for _ in range(6):
        p(doc, "", first=False)
    p(doc, "Выполнил: студент группы 250541", align=WD_ALIGN_PARAGRAPH.RIGHT, first=False)
    p(doc, "Власов Р. Е.", align=WD_ALIGN_PARAGRAPH.RIGHT, first=False)
    p(doc, "Проверил: ____________________", align=WD_ALIGN_PARAGRAPH.RIGHT, first=False)
    for _ in range(6):
        p(doc, "", first=False)
    p(doc, "Минск 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    doc.add_page_break()


def body(doc: Document) -> None:
    h(doc, "1 Цель работы")
    p(doc, "Целью лабораторной работы является разработка и проверка VHDL-модели цифрового устройства, включающего постоянную память ROM, оперативную память RAM с общей шиной данных и управляющий автомат переноса слова из ячейки-источника в ячейку-приемник.")
    p(doc, "Дополнительно требуется подготовить проект в двух САПР: Intel Quartus II и Xilinx Vivado, получить временные диаграммы, RTL-представление, результаты компиляции и синтеза, а также подтвердить работоспособность автоматическим testbench.")

    h(doc, "2 Задание и исходные данные")
    p(doc, "По методическим указаниям необходимо разработать блок, в который входят модули памяти lpm_rom и lpm_ram_io, подключенные к общей восьмиразрядной шине. Требуется продемонстрировать чтение из памяти и выполнить пересылку данных из ROM в RAM через восьмиразрядный регистр.")
    pic(doc, "variant_table.png", "Рисунок 1 - Таблица вариантов лабораторной работы с выделенным вариантом 4")
    table(doc, [
        ["Параметр", "Значение для варианта 4"],
        ["Синхронность ввода lpm_rom", "Асинхронный ввод адреса"],
        ["Синхронность вывода lpm_rom", "Синхронный вывод данных"],
        ["Синхронность ввода lpm_ram_io", "Синхронная запись данных"],
        ["Синхронность вывода lpm_ram_io", "Синхронное чтение данных"],
        ["Ячейка-источник", "ROM[4]"],
        ["Ячейка-приемник", "RAM[5]"],
        ["Передаваемое значение", "0x5A"],
    ], [6.2, 9.4])
    pic(doc, "memory_map.png", "Рисунок 2 - Карта памяти для функционального моделирования")

    h(doc, "3 Теоретические сведения")
    p(doc, "Постоянная память ROM хранит заранее заданный набор слов и в данной работе используется как источник данных. Оперативная память RAM допускает запись и последующее чтение по адресу. В FPGA внутренние двунаправленные шины и высокоимпедансные состояния обычно преобразуются синтезатором в логические мультиплексоры, однако в функциональной модели VHDL состояние 'Z' удобно показывает, какой узел в данный момент управляет общей шиной.")
    p(doc, "Синхронный вывод памяти означает, что новое значение появляется на выходе только после фронта тактового сигнала. Поэтому для варианта 4 чтение ROM и чтение RAM имеют задержку на один такт. Это учитывается управляющим автоматом: после включения ROM выполняется отдельное состояние загрузки регистра, а после записи в RAM выполняется состояние синхронного чтения.")
    p(doc, "Модель lpm_rom реализована как массив из шестнадцати восьмиразрядных слов. Модель lpm_ram_io содержит массив RAM 16 x 8, вход записи, вход разрешения вывода и общий порт data_io. При записи значение берется с общей шины, а при чтении RAM сама выводит слово на эту шину.")
    pic(doc, "memory_architecture.png", "Рисунок 3 - Иерархическое представление разрабатываемого блока")
    pic(doc, "fsm_diagram.png", "Рисунок 4 - Граф состояний управляющего автомата")

    h(doc, "4 Разработка VHDL-модели")
    p(doc, "Главный модуль memory_transfer содержит константы адресов источника и приемника, экземпляры ROM и RAM, регистр данных и конечный автомат. Вариант 4 задан константами SRC_ADDR_C = 4, DST_ADDR_C = 5 и EXPECTED_C = 0x5A.")
    p(doc, "Автомат имеет шесть состояний: S_IDLE, S_READ_ROM, S_LOAD_REG, S_WRITE_RAM, S_READ_RAM и S_DONE. В S_READ_ROM включается ROM, в S_LOAD_REG регистр захватывает слово 0x5A, в S_WRITE_RAM регистр выдает слово на шину и активируется запись RAM, в S_READ_RAM выполняется синхронное чтение RAM, а в S_DONE формируются done_o и verify_ok_o.")
    pic(doc, "project_files_diagram.png", "Рисунок 5 - Структура папки проекта")
    table(doc, [
        ["Сигнал", "Назначение"],
        ["clk_i", "Тактовый сигнал всех синхронных элементов."],
        ["rst_i", "Сброс автомата в состояние ожидания и очистка регистра."],
        ["start_i", "Запуск операции переноса."],
        ["data_bus_o", "Наблюдаемая общая шина данных."],
        ["rom_q_o, reg_q_o, ram_q_o", "Контрольные выходы ROM, регистра и RAM для моделирования."],
        ["done_o", "Операция переноса завершена."],
        ["verify_ok_o", "Прочитанное из RAM значение совпадает с ожидаемым 0x5A."],
    ], [4.8, 10.8])

    h(doc, "5 Функциональное моделирование")
    p(doc, "Testbench tb_memory_transfer формирует тактовый сигнал с периодом 10 ns, снимает reset, подает start и по тактам проверяет состояние автомата, адреса памяти, состояние общей шины, содержимое регистра и флаг verify_ok. Проверки выполнены операторами assert, поэтому ошибка сразу завершила бы моделирование с severity failure.")
    pic(doc, "testbench_coverage.png", "Рисунок 6 - Покрытие проверок testbench")
    pic(doc, "timing_diagram.png", "Рисунок 7 - Полная временная диаграмма переноса данных")
    pic(doc, "timing_bus_zoom.png", "Рисунок 8 - Увеличенный фрагмент работы общей шины")
    pic(doc, "ghdl_result.png", "Рисунок 9 - Результат запуска GHDL")
    p(doc, "По временной диаграмме видно, что ROM начинает выдавать 0x5A после тактового фронта, затем это значение фиксируется в регистре, после чего регистр управляет общей шиной при записи в RAM[5]. На завершающем этапе RAM синхронно выводит 0x5A, и флаг verify_ok_o становится равным единице.")

    h(doc, "6 Реализация проекта в Intel Quartus II")
    p(doc, "Для Quartus подготовлены файлы memory_transfer.qpf и memory_transfer.qsf. Верхним уровнем назначен memory_transfer_rtl_view_top, в котором биты общей шины выведены отдельными портами BUS_D0...BUS_D7 для удобного просмотра RTL-схемы.")
    p(doc, "В методических указаниях указано семейство Flex10K. Установленная версия Quartus II 9.1 Web Edition не принимает строки FAMILY \"FLEX10K\" и FAMILY \"FLEX10KE\", что зафиксировано в логе проверки. Для получения полной компиляции и отчетов проект собран на доступном семействе Stratix II. VHDL-описание при этом не изменялось.")
    pic(doc, "rtl_quartus_view.png", "Рисунок 10 - RTL-представление проекта для Quartus II")
    pic(doc, "quartus_result.png", "Рисунок 11 - Итог компиляции Quartus II")
    p(doc, "Quartus успешно выполнил Analysis & Synthesis, Fitter, Assembler и Classic Timing Analyzer. В отчете синтеза RAM распознана как altsyncram, что подтверждает аппаратную интерпретацию массива памяти, а не только поведенческое моделирование.")

    h(doc, "7 Реализация проекта в Xilinx Vivado")
    p(doc, "Для Vivado подготовлены create_project.tcl, memory_transfer.xdc и скрипт run_vivado_synth.ps1. Запуск выполняется через временную ASCII-папку D:/vivado_stage, потому что Vivado 2022.1 нестабильно работает с кириллицей в пути проекта. После синтеза отчеты и checkpoint-файлы копируются обратно в папку лабораторной работы.")
    pic(doc, "rtl_vivado_view.png", "Рисунок 12 - RTL-представление проекта для Vivado")
    pic(doc, "vivado_result.png", "Рисунок 13 - Итог синтеза Vivado")
    pic(doc, "resource_comparison.png", "Рисунок 14 - Сводка использования ресурсов в двух САПР")
    p(doc, "Vivado успешно синтезировал проект для xc7a35tcpg236-1. В отчете указано 8 LUT as Distributed RAM, что соответствует реализации небольшой RAM 16 x 8 во внутренней распределенной памяти FPGA.")

    h(doc, "8 Контрольная проверка результатов")
    table(doc, [
        ["Проверяемый пункт", "Результат"],
        ["VHDL-синтаксис", "Принят GHDL, Quartus II и Vivado."],
        ["Функциональная модель", "GHDL: TEST PASSED на 66 ns."],
        ["Перенос данных", "ROM[4] = 0x5A записано в RAM[5]."],
        ["Синхронность ROM/RAM", "На временной диаграмме видна задержка вывода на такт."],
        ["Quartus II", "Full Compilation successful, 0 errors."],
        ["Vivado", "synth_design completed successfully, 0 errors."],
        ["Отчет", "Сформированы DOCX/PDF, временные диаграммы, RTL-схемы и изображения результатов."],
    ], [6.2, 9.4])

    h(doc, "9 Вывод")
    p(doc, "В ходе лабораторной работы разработана VHDL-модель блока памяти по варианту 4. Реализованы модели ROM и RAM, общая шина данных, восьмиразрядный регистр и управляющий автомат переноса данных из ROM[4] в RAM[5].")
    p(doc, "Работоспособность подтверждена автоматическим testbench и временными диаграммами GHDL. Проект успешно собран в Quartus II и синтезирован в Xilinx Vivado. Получены RTL-представления, отчеты ресурсов и файлы проектов, поэтому лабораторная работа готова к демонстрации и защите.")

    doc.add_page_break()
    h(doc, "Приложение А. Листинги VHDL-модулей")
    code(doc, "А.1 Модель ROM", ROOT / "src" / "lpm_rom_model.vhd")
    code(doc, "А.2 Модель RAM", ROOT / "src" / "lpm_ram_io_model.vhd")
    code(doc, "А.3 Управляющий модуль memory_transfer", ROOT / "src" / "memory_transfer.vhd")
    code(doc, "А.4 Верхний уровень для RTL Viewer", ROOT / "src" / "memory_transfer_rtl_view_top.vhd")
    doc.add_page_break()
    h(doc, "Приложение Б. Листинг testbench")
    code(doc, "Б.1 tb_memory_transfer", ROOT / "tb" / "tb_memory_transfer.vhd")
    doc.add_page_break()
    h(doc, "Приложение В. Файлы проектов САПР")
    code(doc, "В.1 Quartus QSF", ROOT / "quartus" / "memory_transfer.qsf")
    code(doc, "В.2 Vivado TCL", ROOT / "vivado" / "create_project.tcl")


def main() -> None:
    REPORT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    title_page(doc)
    body(doc)
    out = REPORT / "250541_L2_Власов_РЕ_АПЦУ.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
