from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
IMAGES = ROOT / "images"
SRC = ROOT / "src" / "mod11_counter.vhd"
TB = ROOT / "tb" / "tb_mod11_counter.vhd"
RTL_TOP = ROOT / "src" / "mod11_counter_rtl_view_top.vhd"
QSF = ROOT / "quartus" / "mod11_counter.qsf"
VIVADO_TCL = ROOT / "vivado" / "create_project.tcl"
XDC = ROOT / "vivado" / "mod11_counter.xdc"


def set_run_style(run, size: float = 14, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_format(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, before=0, after=0, line=1.0) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(1.25) if first_line else Cm(0)
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(doc: Document, text: str = "", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, bold=False, size=14, before=0, after=0) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=align, first_line=first_line, before=before, after=after)
    run = p.add_run(text)
    set_run_style(run, size=size, bold=bold)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=True, before=10, after=6)
    run = p.add_run(text)
    set_run_style(run, bold=True)
    p.style = doc.styles["Heading 1"]


def add_appendix_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=10, after=8)
    run = p.add_run(text)
    set_run_style(run, bold=True)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=2, after=8)
    run = p.add_run(text)
    set_run_style(run)


def set_cell(cell, text: str, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=14) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_style(run, size=size, bold=bold)


def add_table(doc: Document, rows: list[list[str]], widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            if widths_cm:
                cell.width = Cm(widths_cm[c])
            set_cell(cell, text, bold=(r == 0), align=WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    add_text(doc, "", first_line=False, after=0)


def add_picture(doc: Document, path: Path, caption: str, width_cm: float = 15.6) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=4, after=2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_style(run, size=12)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_code_block(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, before=0, after=0, line=1.0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Cm(1.25)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.0


def title_page(doc: Document) -> None:
    title_lines = [
        "Министерство образования Республики Беларусь",
        "",
        "Учреждение образования",
        "БЕЛОРУССКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ",
        "ИНФОРМАТИКИ И РАДИОЭЛЕКТРОНИКИ",
        "",
        "Факультет компьютерных систем и сетей",
        "",
        "Кафедра электронных вычислительных машин",
        "",
        "Дисциплина: Автоматизированное проектирование цифровых устройств",
    ]
    for line in title_lines:
        add_text(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    for _ in range(5):
        add_text(doc, "", first_line=False)

    add_text(doc, "ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ №1", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, bold=True)
    add_text(doc, "на тему", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "«Разработка математической и программной модели цифрового объекта", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "“Счетчик с принудительным порядком счета”»", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

    for _ in range(5):
        add_text(doc, "", first_line=False)

    add_text(doc, "Выполнил: студент группы 250541", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    add_text(doc, "Власов Р. Е.", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    add_text(doc, "Проверил: ____________________", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)

    for _ in range(6):
        add_text(doc, "", first_line=False)

    add_text(doc, "Минск 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_page_break()


def contents_page(doc: Document) -> None:
    add_text(doc, "СОДЕРЖАНИЕ", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, bold=True)
    contents = [
        ("1 Цель работы", "3"),
        ("2 Исходные данные к работе", "3"),
        ("3 Теоретические сведения", "4"),
        ("4 Разработка программной модели", "5"),
        ("5 Выполнение работы в САПР", "5"),
        ("6 Моделирование и проверка", "6"),
        ("7 Вывод", "8"),
        ("Приложение А Листинг VHDL-модулей", "9"),
        ("Приложение Б Листинг тестбенча", "11"),
        ("Приложение В Файлы проектов САПР", "13"),
        ("Приложение Г Выдержки из логов проверки", "14"),
    ]
    for title, page in contents:
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
        run = p.add_run(f"{title}\t{page}")
        set_run_style(run)
    doc.add_page_break()


def body(doc: Document) -> None:
    add_heading(doc, "1 Цель работы")
    add_text(
        doc,
        "Целью лабораторной работы является разработка математической и программной модели цифрового объекта "
        "«Счетчик с принудительным порядком счета», описание объекта на языке VHDL, функциональная проверка "
        "модели и подготовка проекта для двух САПР: Intel Quartus II и Xilinx Vivado.",
    )
    add_text(
        doc,
        "В ходе выполнения работы необходимо определить последовательность состояний счетчика, разработать "
        "VHDL-описание и тестбенч, выполнить моделирование, получить временную диаграмму, подготовить RTL-представление "
        "и зафиксировать результаты анализа проекта.",
    )
    add_caption(doc, "Таблица 1 - Этапы выполнения лабораторной работы")
    add_table(
        doc,
        [
            ["Этап", "Содержание"],
            ["1", "Формирование математической модели счетчика по модулю 11."],
            ["2", "Разработка VHDL-модуля и верхнего уровня для просмотра RTL-схемы."],
            ["3", "Функциональное моделирование и построение временной диаграммы."],
            ["4", "Подготовка проектов Quartus II и Xilinx Vivado."],
            ["5", "Анализ результатов синтеза и оформление отчета."],
        ],
        [2.0, 13.8],
    )

    add_heading(doc, "2 Исходные данные к работе")
    add_text(
        doc,
        "По заданию требуется разработать счетчик по модулю 11 с принудительной последовательностью счета "
        "5-15, 5-15. Таким образом, рабочими являются одиннадцать состояний четырехразрядного двоичного счетчика: "
        "0101, 0110, ..., 1111. После состояния 1111 счетчик принудительно возвращается в состояние 0101.",
    )
    add_caption(doc, "Таблица 2 - Исходные данные")
    add_table(
        doc,
        [
            ["Параметр", "Значение"],
            ["Тип объекта", "Синхронный двоичный счетчик с принудительным порядком счета"],
            ["Модуль счета", "11"],
            ["Последовательность", "5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 5"],
            ["Разрядность выхода", "4 бита"],
            ["Управление", "rst - сброс к 5; load - загрузка 5; en - разрешение счета"],
            ["Выход конца счета", "tc=1 при q=15 и en=1"],
        ],
        [5.0, 10.8],
    )
    add_caption(doc, "Таблица 3 - Назначение сигналов VHDL-модуля")
    add_table(
        doc,
        [
            ["Сигнал", "Направление", "Назначение"],
            ["clk_i", "in", "Тактовый сигнал."],
            ["rst_i", "in", "Синхронный сброс счетчика в состояние 5."],
            ["en_i", "in", "Разрешение перехода к следующему состоянию."],
            ["load_i", "in", "Принудительная загрузка начального состояния 5."],
            ["q_o[3:0]", "out", "Текущее четырехразрядное состояние счетчика."],
            ["tc_o", "out", "Признак достижения конечного состояния при активном разрешении счета."],
        ],
        [3.2, 3.2, 9.4],
    )

    add_heading(doc, "3 Теоретические сведения")
    add_text(
        doc,
        "Счетчик является последовательностным цифровым устройством, состояние которого изменяется под действием "
        "тактового сигнала. В обычном двоичном счетчике состояния следуют в естественном порядке, однако в счетчике "
        "с принудительным порядком часть состояний исключается, а переходы задаются условиями управления.",
    )
    add_text(
        doc,
        "Для рассматриваемого объекта используются состояния от 5 до 15. Состояния от 0 до 4 не входят в рабочий "
        "цикл. Поэтому при сбросе, загрузке начального значения и при переходе из конечного состояния устройство "
        "должно принудительно устанавливать код 0101.",
    )
    add_text(
        doc,
        "Математическая модель переходов имеет вид:",
    )
    add_text(
        doc,
        "q(k+1)=5 при rst=1 или load=1; q(k+1)=5 при en=1 и q(k)=15; q(k+1)=q(k)+1 при en=1 и 5<=q(k)<15; q(k+1)=q(k) при en=0.",
        first_line=False,
    )
    add_picture(doc, IMAGES / "state_diagram.png", "Рисунок 1 - Граф состояний счетчика по модулю 11", 15.8)

    add_heading(doc, "4 Разработка программной модели")
    add_text(
        doc,
        "Программная модель разработана на языке VHDL с использованием библиотек ieee.std_logic_1164 и "
        "ieee.numeric_std. Внутреннее состояние хранится в регистре count_reg типа unsigned(3 downto 0). "
        "Минимальное и максимальное состояния вынесены в константы COUNT_MIN_C и COUNT_MAX_C.",
    )
    add_text(
        doc,
        "Основной фрагмент VHDL-описания приведен ниже. Сначала обрабатываются сигналы rst_i и load_i, затем "
        "при активном en_i выполняется проверка конечного состояния и переход к следующему коду.",
    )
    add_code_block(
        doc,
        "\n".join(
            [
                "process(clk_i)",
                "begin",
                "  if rising_edge(clk_i) then",
                "    if rst_i = '1' or load_i = '1' then",
                "      count_reg <= COUNT_MIN_C;",
                "    elsif en_i = '1' then",
                "      if count_reg = COUNT_MAX_C then",
                "        count_reg <= COUNT_MIN_C;",
                "      else",
                "        count_reg <= count_reg + 1;",
                "      end if;",
                "    end if;",
                "  end if;",
                "end process;",
            ]
        ),
    )
    add_text(
        doc,
        "Для просмотра иерархии в RTL Viewer дополнительно создан верхний уровень mod11_counter_rtl_view_top. "
        "Он выводит четыре разряда счетчика отдельными портами Q0-Q3 и содержит экземпляр основного модуля U_COUNTER.",
    )

    add_heading(doc, "5 Выполнение работы в САПР")
    add_text(
        doc,
        "Для Intel Quartus II создан проект mod11_counter.qpf с файлом назначений mod11_counter.qsf. В качестве "
        "семейства ПЛИС выбрано Stratix II, верхним уровнем назначен mod11_counter_rtl_view_top. Компиляция выполнена "
        "в Quartus II 9.1 через скрипт run_quartus_compile.ps1.",
    )
    add_picture(doc, IMAGES / "rtl_quartus_view.png", "Рисунок 2 - RTL-представление проекта в Quartus II", 15.8)
    add_text(
        doc,
        "Для Xilinx Vivado подготовлен TCL-скрипт create_project.tcl, который создает проект, добавляет VHDL-файлы, "
        "подключает тестбенч и шаблон ограничений XDC. Синтез проекта проверен в Vivado ML Standard 2022.1. Так как "
        "эта версия Vivado нестабильно работает с кириллицей в пути проекта, запуск выполнен через ASCII-каталог "
        "D:/vivado_stage с последующим копированием результатов в папку лабораторной работы.",
    )
    add_picture(doc, IMAGES / "rtl_vivado_view.png", "Рисунок 3 - RTL-представление проекта для Xilinx Vivado", 15.8)
    add_caption(doc, "Таблица 4 - Состав проектных файлов")
    add_table(
        doc,
        [
            ["Файл", "Назначение"],
            ["src/mod11_counter.vhd", "Основной VHDL-модуль счетчика."],
            ["src/mod11_counter_rtl_view_top.vhd", "Верхний уровень для RTL Viewer."],
            ["tb/tb_mod11_counter.vhd", "Тестбенч для функциональной проверки."],
            ["quartus/mod11_counter.qsf", "Назначения проекта Quartus II."],
            ["vivado/create_project.tcl", "Создание проекта Vivado и запуск синтеза."],
            ["vivado/mod11_counter.xdc", "Шаблон временных и физических ограничений Vivado."],
        ],
        [6.2, 9.6],
    )

    add_heading(doc, "6 Моделирование и проверка")
    add_text(
        doc,
        "Функциональная проверка выполнена в GHDL. Тестбенч проверяет начальное состояние после сброса, счет "
        "от 5 до 15, формирование признака tc, возврат из 15 в 5, удержание состояния при en=0 и принудительную "
        "загрузку начального состояния сигналом load.",
    )
    add_picture(doc, IMAGES / "timing_diagram.png", "Рисунок 4 - Временная диаграмма работы счетчика", 15.8)
    add_caption(doc, "Таблица 5 - Результаты функционального моделирования")
    add_table(
        doc,
        [
            ["Проверка", "Ожидаемый результат", "Итог"],
            ["Сброс rst", "q=5", "Выполнено"],
            ["Счет при en=1", "q=5,6,...,15", "Выполнено"],
            ["Переход из 15", "q=5 на следующем такте", "Выполнено"],
            ["Удержание при en=0", "q не изменяется", "Выполнено"],
            ["Загрузка load", "q=5", "Выполнено"],
        ],
        [5.2, 6.5, 4.1],
    )
    add_text(
        doc,
        "По результатам запуска тестбенча получено сообщение: tb_mod11_counter: TEST PASSED. Моделирование "
        "остановлено на отметке 180 ns, файл временной диаграммы сохранен как sim/mod11_counter.vcd.",
    )
    add_caption(doc, "Таблица 6 - Результаты проверки в САПР")
    add_table(
        doc,
        [
            ["Показатель", "Значение"],
            ["Версия Quartus II", "9.1 Build 222 Web Edition"],
            ["Статус Analysis & Synthesis", "Successful, 0 errors, 0 warnings"],
            ["Статус Full Compilation", "Successful, 0 errors, 5 warnings"],
            ["Семейство / устройство", "Stratix II / EP2S15F484C3"],
            ["Combinational ALUTs", "9 / 12,480 (<1%) после Fitter"],
            ["Dedicated logic registers", "7 / 12,480 (<1%) после Fitter"],
            ["Total pins", "9 / 343 (3%)"],
            ["Clock Setup", "до 500.00 MHz, failed paths: 0"],
            ["Версия Vivado", "Vivado ML Standard 2022.1"],
            ["Статус Vivado synthesis", "synth_design completed successfully, 0 errors"],
            ["Ресурсы Vivado", "Slice LUTs: 4; Slice Registers: 4; Bonded IOB: 9; BUFG: 1"],
        ],
        [6.2, 9.6],
    )
    add_text(
        doc,
        "Предупреждения Quartus связаны с отсутствием точных назначений физических выводов и временных ограничений "
        "для учебного проекта. Они не являются ошибками VHDL-модели и не нарушают функциональную проверку счетчика.",
    )

    add_heading(doc, "7 Вывод")
    add_text(
        doc,
        "В ходе лабораторной работы разработана математическая модель счетчика по модулю 11 с принудительным "
        "порядком счета 5-15. На языке VHDL создан синхронный счетчик с сигналами сброса, разрешения счета, "
        "принудительной загрузки и признаком конечного состояния.",
    )
    add_text(
        doc,
        "Функциональное моделирование подтвердило корректность последовательности состояний и управляющих режимов. "
        "Проект Quartus II успешно скомпилирован без ошибок, получены RTL-представление, временная диаграмма и "
        "сводка использования ресурсов. Для второй САПР подготовлен и проверен проект Xilinx Vivado на основе того же VHDL-описания.",
    )

    add_appendices(doc)


def add_file_listing(doc: Document, title: str, path: Path) -> None:
    add_text(doc, title, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, bold=True, before=4, after=3)
    add_code_block(doc, path.read_text(encoding="utf-8"))


def add_log_excerpt(doc: Document, title: str, lines: list[str]) -> None:
    add_text(doc, title, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, bold=True, before=4, after=3)
    add_code_block(doc, "\n".join(lines))


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    add_appendix_heading(doc, "ПРИЛОЖЕНИЕ А")
    add_appendix_heading(doc, "Листинг VHDL-модулей")
    add_text(
        doc,
        "В приложении приведены полные тексты VHDL-модулей, использованных при синтезе и просмотре RTL-структуры.",
    )
    add_file_listing(doc, "А.1 Основной модуль mod11_counter.vhd", SRC)
    add_file_listing(doc, "А.2 Верхний уровень mod11_counter_rtl_view_top.vhd", RTL_TOP)

    doc.add_page_break()
    add_appendix_heading(doc, "ПРИЛОЖЕНИЕ Б")
    add_appendix_heading(doc, "Листинг тестбенча")
    add_text(
        doc,
        "Тестбенч выполняет автоматизированную проверку всех основных режимов счетчика. Наличие assert с severity failure "
        "позволяет завершать моделирование ошибкой при нарушении ожидаемой последовательности.",
    )
    add_file_listing(doc, "Б.1 Тестбенч tb_mod11_counter.vhd", TB)

    doc.add_page_break()
    add_appendix_heading(doc, "ПРИЛОЖЕНИЕ В")
    add_appendix_heading(doc, "Файлы проектов САПР")
    add_text(
        doc,
        "Ниже приведены ключевые проектные файлы, позволяющие повторить синтез в Quartus II и Vivado.",
    )
    add_file_listing(doc, "В.1 Файл назначений Quartus II mod11_counter.qsf", QSF)
    add_file_listing(doc, "В.2 Скрипт создания проекта Vivado create_project.tcl", VIVADO_TCL)
    add_file_listing(doc, "В.3 Шаблон ограничений Vivado mod11_counter.xdc", XDC)

    doc.add_page_break()
    add_appendix_heading(doc, "ПРИЛОЖЕНИЕ Г")
    add_appendix_heading(doc, "Выдержки из логов проверки")
    add_text(
        doc,
        "В приложении приведены ключевые строки логов, подтверждающие успешное моделирование и синтез проекта.",
    )
    add_log_excerpt(
        doc,
        "Г.1 Результат GHDL",
        [
            "tb_mod11_counter.vhd:77:5:@156ns:(assertion note): tb_mod11_counter: TEST PASSED",
            "ghdl:info: simulation stopped by --stop-time @180ns",
        ],
    )
    add_log_excerpt(
        doc,
        "Г.2 Результат Quartus II",
        [
            "Quartus II Analysis & Synthesis was successful. 0 errors, 0 warnings",
            "Quartus II Fitter was successful. 0 errors, 4 warnings",
            "Quartus II Assembler was successful. 0 errors, 0 warnings",
            "Quartus II Full Compilation was successful. 0 errors, 5 warnings",
            "Quartus II Shell was successful. 0 errors, 5 warnings",
        ],
    )
    add_log_excerpt(
        doc,
        "Г.3 Результат Vivado",
        [
            "Vivado v2022.1 (64-bit)",
            "Got license for feature 'Synthesis' and/or device 'xc7a35t'",
            "Synthesis finished with 0 errors, 0 critical warnings and 1 warnings.",
            "16 Infos, 1 Warnings, 0 Critical Warnings and 0 Errors encountered.",
            "synth_design completed successfully",
            "Slice LUTs: 4; Slice Registers: 4; Bonded IOB: 9; BUFGCTRL: 1",
        ],
    )


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=True, before=6, after=4)
    run = p.add_run(text)
    set_run_style(run, bold=True)


def body(doc: Document) -> None:
    add_heading(doc, "1 Цель работы")
    add_text(
        doc,
        "Целью лабораторной работы является разработка математической и программной модели цифрового объекта "
        "«Счетчик с принудительным порядком счета», описание устройства на языке VHDL, проверка работоспособности "
        "модели при функциональном моделировании и подготовка проекта в двух системах автоматизированного "
        "проектирования: Intel Quartus II и Xilinx Vivado.",
    )
    add_text(
        doc,
        "В результате работы необходимо получить не только текст VHDL-программы, но и наглядные подтверждения "
        "корректности объекта: граф состояний, таблицу переходов, RTL-представление, временные диаграммы, "
        "результаты компиляции и синтеза, а также выдержки из логов проверки.",
    )

    add_heading(doc, "2 Задание и исходные данные")
    add_text(
        doc,
        "По индивидуальному заданию требуется разработать счетчик по модулю 11 с последовательностью счета "
        "5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, после чего счетчик должен возвращаться в состояние 5. "
        "Таким образом, рабочий цикл содержит одиннадцать состояний, а состояния 0-4 не входят в нормальный "
        "порядок счета.",
    )
    add_caption(doc, "Таблица 1 - Исходные данные лабораторной работы")
    add_table(
        doc,
        [
            ["Параметр", "Значение"],
            ["Тип объекта", "Синхронный двоичный счетчик с принудительным порядком счета"],
            ["Модуль счета", "11"],
            ["Рабочая последовательность", "5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 5"],
            ["Разрядность состояния", "4 бита"],
            ["Сигналы управления", "rst_i - сброс; en_i - разрешение счета; load_i - принудительная загрузка"],
            ["Выходные сигналы", "q_o[3:0] - состояние счетчика; tc_o - признак конечного состояния"],
            ["Средства проверки", "GHDL, Intel Quartus II 9.1, Xilinx Vivado 2022.1"],
        ],
        [5.5, 10.3],
    )
    add_picture(doc, IMAGES / "project_files_diagram.png", "Рисунок 1 - Структура файлов лабораторной работы", 15.6)

    add_heading(doc, "3 Теоретические сведения")
    add_text(
        doc,
        "Счетчик относится к последовательностным цифровым устройствам. Его выходное значение зависит не только "
        "от текущих входных сигналов, но и от ранее сохраненного состояния. В синхронном счетчике изменение "
        "состояния происходит только по активному фронту тактового сигнала, поэтому устройство удобно описывать "
        "как конечный автомат с регистром состояния и комбинационной логикой следующего состояния.",
    )
    add_text(
        doc,
        "Обычный двоичный счетчик проходит все коды выбранной разрядности. В данной работе используется счетчик "
        "с принудительным порядком счета: часть двоичных кодов исключена, а после конечного состояния 15 "
        "выполняется принудительный переход к начальному состоянию 5. Такой подход часто используется при "
        "построении делителей частоты, управляющих автоматов, адресных генераторов и циклических устройств, "
        "когда требуется пройти не весь диапазон двоичных чисел, а только заданный набор состояний.",
    )
    add_text(
        doc,
        "Математически переходы счетчика можно описать следующим образом: при активном rst_i или load_i "
        "следующее состояние равно 5; если разрешение счета en_i неактивно, состояние сохраняется; если en_i "
        "активен и текущее состояние равно 15, следующее состояние равно 5; во всех остальных рабочих состояниях "
        "счетчик увеличивается на единицу.",
    )
    add_picture(doc, IMAGES / "state_diagram.png", "Рисунок 2 - Граф состояний счетчика по модулю 11", 15.6)
    add_picture(doc, IMAGES / "transition_table.png", "Рисунок 3 - Таблица переходов счетчика", 15.6)
    add_picture(doc, IMAGES / "algorithm_flowchart.png", "Рисунок 4 - Алгоритм формирования следующего состояния", 14.8)
    add_text(
        doc,
        "Признак конечного состояния tc_o формируется как комбинационная функция текущего состояния и разрешения "
        "счета. Он равен единице только тогда, когда счетчик находится в состоянии 15 и сигнал en_i активен. "
        "Это позволяет внешней схеме определить момент окончания текущего цикла счета.",
    )

    add_heading(doc, "4 Разработка VHDL-модели")
    add_text(
        doc,
        "VHDL-модель построена на базе библиотек ieee.std_logic_1164 и ieee.numeric_std. Для арифметики выбран "
        "тип unsigned, так как состояние счетчика является четырехразрядным беззнаковым числом. Начальное и "
        "конечное значения вынесены в константы COUNT_MIN_C и COUNT_MAX_C, что делает описание более читаемым "
        "и уменьшает вероятность ошибки при изменении диапазона счета.",
    )
    add_picture(doc, IMAGES / "entity_symbol.png", "Рисунок 5 - Интерфейс VHDL-модуля счетчика", 15.2)
    add_caption(doc, "Таблица 2 - Назначение сигналов основного модуля")
    add_table(
        doc,
        [
            ["Сигнал", "Направление", "Назначение"],
            ["clk_i", "in", "Тактовый сигнал, по фронту которого изменяется состояние."],
            ["rst_i", "in", "Синхронный сброс счетчика в состояние 5."],
            ["en_i", "in", "Разрешение перехода к следующему состоянию."],
            ["load_i", "in", "Принудительная загрузка начального состояния 5."],
            ["q_o[3:0]", "out", "Текущее состояние счетчика в двоичном коде."],
            ["tc_o", "out", "Признак достижения конечного состояния 15 при активном en_i."],
        ],
        [3.1, 3.0, 9.7],
    )
    add_subheading(doc, "Фрагмент описания счетчика")
    add_code_block(
        doc,
        "\n".join(
            [
                "process(clk_i)",
                "begin",
                "  if rising_edge(clk_i) then",
                "    if rst_i = '1' or load_i = '1' then",
                "      count_reg <= COUNT_MIN_C;",
                "    elsif en_i = '1' then",
                "      if count_reg = COUNT_MAX_C then",
                "        count_reg <= COUNT_MIN_C;",
                "      else",
                "        count_reg <= count_reg + 1;",
                "      end if;",
                "    end if;",
                "  end if;",
                "end process;",
                "",
                "q_o  <= std_logic_vector(count_reg);",
                "tc_o <= '1' when count_reg = COUNT_MAX_C and en_i = '1' else '0';",
            ]
        ),
    )
    add_text(
        doc,
        "Для просмотра структуры в RTL Viewer создан дополнительный верхний уровень mod11_counter_rtl_view_top. "
        "Он не меняет логику счетчика, а только выводит разряды q_o отдельными портами Q0-Q3 и задает удобные "
        "имена внешних сигналов CLOCK, RST, EN, LOAD и TC.",
    )

    add_heading(doc, "5 Функциональное моделирование")
    add_text(
        doc,
        "Работоспособность VHDL-модели проверена тестбенчем tb_mod11_counter. Тестбенч формирует тактовый сигнал "
        "с периодом 10 ns, подает управляющие воздействия и с помощью assert проверяет ожидаемое значение выхода "
        "q_o после каждого значимого такта. При любой ошибке моделирование завершается с сообщением failure.",
    )
    add_picture(doc, IMAGES / "testbench_coverage.png", "Рисунок 6 - Покрытие проверок в тестбенче", 15.6)
    add_picture(doc, IMAGES / "timing_diagram.png", "Рисунок 7 - Полная временная диаграмма работы счетчика", 15.6)
    add_picture(doc, IMAGES / "timing_wrap_zoom.png", "Рисунок 8 - Увеличенный фрагмент перехода 15 -> 5", 15.6)
    add_picture(doc, IMAGES / "timing_hold_load_zoom.png", "Рисунок 9 - Увеличенный фрагмент удержания и загрузки", 15.6)
    add_picture(doc, IMAGES / "ghdl_result.png", "Рисунок 10 - Результат запуска функционального моделирования GHDL", 15.6)
    add_text(
        doc,
        "По результатам моделирования получено сообщение tb_mod11_counter: TEST PASSED. Это означает, что все "
        "проверяемые режимы счетчика совпали с ожидаемой математической моделью: сброс, счет 5-15, возврат из "
        "15 в 5, удержание при en_i=0, принудительная загрузка load_i и формирование признака tc_o.",
    )

    add_heading(doc, "6 Выполнение проекта в Intel Quartus II")
    add_text(
        doc,
        "Для первой САПР подготовлены файлы проекта mod11_counter.qpf и mod11_counter.qsf. Верхним уровнем "
        "назначен mod11_counter_rtl_view_top, благодаря чему в RTL Viewer виден экземпляр U_COUNTER и внешние "
        "порты, соответствующие учебной схеме. Компиляция выполнена в Quartus II 9.1 Web Edition.",
    )
    add_picture(doc, IMAGES / "rtl_quartus_view.png", "Рисунок 11 - RTL-представление проекта в Quartus II", 15.6)
    add_picture(doc, IMAGES / "quartus_result.png", "Рисунок 12 - Итог компиляции проекта в Quartus II", 15.6)
    add_text(
        doc,
        "Full Compilation завершена успешно. Предупреждения Quartus связаны с тем, что для учебного проекта "
        "не заданы реальные номера выводов платы и полноценные временные ограничения. Эти предупреждения не "
        "являются ошибками VHDL-описания и не влияют на функциональную проверку счетчика.",
    )

    add_heading(doc, "7 Выполнение проекта в Xilinx Vivado")
    add_text(
        doc,
        "Для второй САПР подготовлен TCL-скрипт create_project.tcl. Он создает проект Vivado, добавляет VHDL-файлы, "
        "подключает testbench и файл ограничений mod11_counter.xdc, назначает верхний уровень и запускает синтез. "
        "На данной машине установлен Vivado ML Standard 2022.1.",
    )
    add_text(
        doc,
        "Так как Vivado 2022.1 нестабильно работает с кириллицей в пути проекта, запуск выполняется через "
        "служебный ASCII-каталог D:/vivado_stage. После завершения синтеза отчеты и checkpoint-файлы копируются "
        "обратно в папку лабораторной работы. Такой запуск проверяет тот же самый VHDL-код, но исключает проблему "
        "с кодировкой пути.",
    )
    add_picture(doc, IMAGES / "rtl_vivado_view.png", "Рисунок 13 - RTL-представление проекта для Xilinx Vivado", 15.6)
    add_picture(doc, IMAGES / "vivado_result.png", "Рисунок 14 - Итог синтеза проекта в Vivado", 15.6)
    add_picture(doc, IMAGES / "resource_comparison.png", "Рисунок 15 - Сводка использования ресурсов в двух САПР", 15.6)

    doc.add_page_break()
    add_heading(doc, "8 Контрольная проверка результатов")
    add_caption(doc, "Таблица 3 - Итоговая проверка лабораторной работы")
    add_table(
        doc,
        [
            ["Проверяемый пункт", "Результат"],
            ["VHDL-код основного счетчика", "Разработан, синтаксис принят GHDL, Quartus II и Vivado."],
            ["Последовательность счета", "Подтверждена моделированием: 5, 6, ..., 15, 5."],
            ["Удержание при en_i=0", "Проверено тестбенчем, значение q_o не изменяется."],
            ["Сброс и загрузка", "rst_i и load_i принудительно устанавливают состояние 5."],
            ["Признак tc_o", "Активен при q_o=15 и en_i=1."],
            ["Quartus II", "Full Compilation successful, 0 errors."],
            ["Vivado", "synth_design completed successfully, 0 errors."],
            ["Отчет", "DOCX/PDF сформирован, страницы отрендерены для визуальной проверки."],
        ],
        [6.4, 9.4],
    )
    add_text(
        doc,
        "Таким образом, счетчик является рабочим и адекватным заданию: его функциональная модель проходит "
        "автоматические проверки, а синтез в двух САПР подтверждает, что описание пригодно для аппаратной "
        "реализации в ПЛИС.",
    )

    add_heading(doc, "9 Вывод")
    add_text(
        doc,
        "В ходе лабораторной работы разработана математическая модель счетчика по модулю 11 с принудительным "
        "порядком счета 5-15. На языке VHDL создан синхронный счетчик с управляющими сигналами rst_i, en_i, "
        "load_i и выходным признаком tc_o.",
    )
    add_text(
        doc,
        "Работа счетчика подтверждена функциональным моделированием в GHDL, временными диаграммами и проверками "
        "testbench. Проект успешно обработан в Intel Quartus II и Xilinx Vivado, получены RTL-представления, "
        "результаты компиляции, результаты синтеза и сводки использования ресурсов. Подготовленный комплект "
        "может быть использован для защиты первой лабораторной работы.",
    )

    add_appendices(doc)


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    title_page(doc)
    body(doc)
    out = REPORT_DIR / "250541_L1_Власов_РЕ_АПЦУ.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
