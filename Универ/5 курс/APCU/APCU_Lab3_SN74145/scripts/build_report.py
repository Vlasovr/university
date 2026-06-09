from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
IMG = ROOT / "images"


def set_run(run, size=14, bold=False, italic=False, mono=False):
    name = "Courier New" if mono else "Times New Roman"
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def fmt(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=True, before=0, after=0, line=1.0):
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(1.25) if first else Cm(0)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def p(doc, text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=True, bold=False, italic=False, size=14, before=0, after=0):
    par = doc.add_paragraph()
    fmt(par, align=align, first=first, before=before, after=after)
    run = par.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)


def h(doc, text):
    par = doc.add_paragraph()
    par.style = doc.styles["Heading 1"]
    fmt(par, align=WD_ALIGN_PARAGRAPH.LEFT, first=True, before=10, after=6)
    run = par.add_run(text)
    set_run(run, bold=True)


def sh(doc, text):
    par = doc.add_paragraph()
    par.style = doc.styles["Heading 2"]
    fmt(par, align=WD_ALIGN_PARAGRAPH.LEFT, first=True, before=8, after=4)
    run = par.add_run(text)
    set_run(run, bold=True)


def caption(doc, text):
    par = doc.add_paragraph()
    fmt(par, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, before=2, after=7)
    run = par.add_run(text)
    set_run(run)


def pic(doc, filename, text, width=15.6):
    par = doc.add_paragraph()
    fmt(par, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, before=4, after=2)
    par.add_run().add_picture(str(IMG / filename), width=Cm(width))
    caption(doc, text)


def cell(cell_obj, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=14):
    cell_obj.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    par = cell_obj.paragraphs[0]
    fmt(par, align=align, first=False, before=0, after=0)
    run = par.add_run(text)
    set_run(run, size=size, bold=bold)


def table(doc, rows, widths=None):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl.autofit = False
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cl = tbl.cell(r, c)
            if widths:
                cl.width = Cm(widths[c])
            cell(cl, text, bold=(r == 0), align=WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    p(doc, "", first=False)


def code(doc, title, path):
    sh(doc, title)
    for line in path.read_text(encoding="utf-8").splitlines():
        par = doc.add_paragraph()
        fmt(par, align=WD_ALIGN_PARAGRAPH.LEFT, first=False, before=0, after=0, line=1.0)
        run = par.add_run(line)
        set_run(run, size=8.5, mono=True)


def page_number(section):
    par = section.footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    set_run(run, size=12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)
    page_number(sec)
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.color.rgb = RGBColor(0, 0, 0)
        if name.startswith("Heading"):
            style.font.bold = True


def title_page(doc):
    for line in [
        "Министерство образования Республики Беларусь",
        "",
        "Учреждение образования",
        "БЕЛОРУССКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ",
        "ИНФОРМАТИКИ И РАДИОЭЛЕКТРОНИКИ",
        "",
        "Факультет компьютерных систем и сетей",
        "",
        "Кафедра электронных вычислительных машин",
        "",
        "Дисциплина: Автоматизированное проектирование цифровых устройств",
    ]:
        p(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    for _ in range(5):
        p(doc, "", first=False)
    p(doc, "ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ №1", align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True)
    p(doc, "на тему", align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    p(doc, "«Проектирование комбинационных схем с использованием языка VHDL»", align=WD_ALIGN_PARAGRAPH.CENTER, first=False, bold=True)
    p(doc, "Вариант 4. Микросхема 145 BCD-to-decimal decoder/driver", align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    for _ in range(6):
        p(doc, "", first=False)
    p(doc, "Выполнил: студент группы 250541", align=WD_ALIGN_PARAGRAPH.RIGHT, first=False)
    p(doc, "Власов Р. Е.", align=WD_ALIGN_PARAGRAPH.RIGHT, first=False)
    p(doc, "Проверил: ____________________", align=WD_ALIGN_PARAGRAPH.RIGHT, first=False)
    for _ in range(6):
        p(doc, "", first=False)
    p(doc, "Минск 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    doc.add_page_break()


def body(doc):
    h(doc, "1 Цель работы")
    p(doc, "Целью лабораторной работы является ознакомление с САПР Xilinx Vivado, изучение базового синтаксиса VHDL и приобретение навыков проектирования простейших комбинационных схем на языке VHDL.")
    p(doc, "В соответствии с вариантом 4 необходимо реализовать комбинационное устройство по схеме из файла задания 4.pdf. Устройство должно быть описано двумя способами: через логические операторы и параллельные безусловные присваивания, а также через последовательные операторы. Для проверки требуется полный перебор входных значений.")

    h(doc, "2 Исходное задание")
    p(doc, "Варианту 4 соответствует микросхема 145 BCD-to-decimal decoders/drivers. Это BCD-десятичный дешифратор с четырьмя входами A, B, C, D и десятью активными-низкими выходами OUTPUT 0...OUTPUT 9.")
    pic(doc, "task_variant4_datasheet.png", "Рисунок 1 - Исходная схема и таблица функционирования из файла 4.pdf", 13.2)
    pic(doc, "truth_table.png", "Рисунок 2 - Таблица истинности, использованная для VHDL-модели")
    table(doc, [
        ["Порт", "Назначение"],
        ["A", "Младший разряд BCD-кода."],
        ["B", "Второй разряд BCD-кода."],
        ["C", "Третий разряд BCD-кода."],
        ["D", "Старший разряд BCD-кода."],
        ["Y0...Y9", "Десятичные выходы, активные низким уровнем."],
    ], [4.0, 11.6])

    h(doc, "3 Анализ схемы и принцип работы")
    p(doc, "Микросхема 145 преобразует четырехразрядный BCD-код в десятичный позиционный код. Для входных кодов от 0000 до 1001 активируется ровно один выход, соответствующий числу на входе. Активный уровень выхода - логический ноль. Остальные выходы находятся в логической единице.")
    p(doc, "Для невалидных BCD-кодов 1010...1111 ни один десятичный выход не должен активироваться, поэтому все выходы остаются в единице. Такое поведение видно в таблице функционирования исходного задания.")
    p(doc, "Комбинационный характер схемы означает, что выходы зависят только от текущего набора входов A, B, C, D. В корректном результате синтеза не должно быть регистров или защелок.")
    pic(doc, "entity_symbol.png", "Рисунок 3 - Условное обозначение разработанного VHDL-модуля")
    pic(doc, "logic_equations.png", "Рисунок 4 - Логические выражения для активных-низких выходов")

    h(doc, "4 Разработка VHDL-описаний")
    p(doc, "Первая реализация sn74145_concurrent построена только на логических операторах not/or и параллельных операторах безусловного присваивания. Например, выход Y0 равен A or B or C or D, поэтому он становится нулем только при входном коде 0000. Аналогично выход Y1 становится нулем только при коде 0001, выход Y2 - при 0010 и так далее.")
    p(doc, "Вторая реализация sn74145_sequential использует процесс и оператор case. Сначала все десять выходов устанавливаются в '1', затем для входных кодов 0...9 соответствующий выход устанавливается в '0'. Для остальных кодов ветка others ничего не меняет, поэтому все выходы остаются неактивными.")
    p(doc, "Для просмотра структуры в RTL Viewer создан верхний уровень sn74145_compare_top. Он подключает обе реализации к одним и тем же входам и формирует сигнал MATCH, равный единице при совпадении двух выходных векторов.")
    pic(doc, "compare_topology.png", "Рисунок 5 - Верхний уровень сравнения двух реализаций")
    pic(doc, "project_files_diagram.png", "Рисунок 6 - Структура папки проекта")

    h(doc, "5 Функциональное моделирование")
    p(doc, "Testbench tb_sn74145 выполняет полный перебор всех 16 входных комбинаций A, B, C, D. Для каждого кода формируется ожидаемый десятиразрядный вектор: для кодов 0...9 один выход равен нулю, для кодов 10...15 все выходы равны единице.")
    pic(doc, "testbench_coverage.png", "Рисунок 7 - Покрытие проверок testbench")
    pic(doc, "timing_diagram.png", "Рисунок 8 - Временная диаграмма полного перебора входов")
    pic(doc, "ghdl_result.png", "Рисунок 9 - Результат запуска GHDL")
    p(doc, "Симуляция завершилась сообщением TEST PASSED на 160 ns. Это подтверждает совпадение обеих реализаций с таблицей истинности и между собой.")

    h(doc, "6 Реализация проекта в Intel Quartus II")
    p(doc, "Для Quartus II подготовлены файлы sn74145.qpf и sn74145.qsf. Верхним уровнем назначен sn74145_compare_top, что позволяет увидеть обе реализации и сигнал сравнения в RTL-представлении.")
    pic(doc, "rtl_quartus_view.png", "Рисунок 10 - RTL-представление проекта в Quartus II")
    pic(doc, "quartus_result.png", "Рисунок 11 - Итог компиляции Quartus II")
    p(doc, "Полная компиляция Quartus II завершена без ошибок. Использовано 10 комбинационных ALUT и 0 регистров. Предупреждение о том, что MATCH stuck at VCC, является ожидаемым: обе реализации эквивалентны, поэтому сигнал сравнения всегда равен единице.")

    h(doc, "7 Реализация проекта в Xilinx Vivado")
    p(doc, "Для Vivado подготовлены create_project.tcl, sn74145.xdc и run_vivado_synth.ps1. Синтез выполняется через ASCII-папку D:/vivado_stage, после чего отчеты и checkpoint-файлы копируются обратно в папку лабораторной работы.")
    pic(doc, "rtl_vivado_view.png", "Рисунок 12 - RTL-представление проекта в Vivado")
    pic(doc, "vivado_result.png", "Рисунок 13 - Итог синтеза Vivado")
    pic(doc, "resource_comparison.png", "Рисунок 14 - Сводка ресурсов в двух САПР")
    p(doc, "Vivado успешно синтезировал проект для xc7a35tcpg236-1. В отчете указано 9 Slice LUTs и 0 Slice Registers, что соответствует комбинационной схеме.")

    h(doc, "8 Контрольная проверка")
    table(doc, [
        ["Проверяемый пункт", "Результат"],
        ["Вариант задания", "4.pdf: 145 BCD-to-decimal decoders/drivers."],
        ["Первая реализация", "Только логические операторы и параллельные безусловные присваивания."],
        ["Вторая реализация", "Последовательное описание process/case."],
        ["Полный перебор", "Проверены все 16 входных кодов."],
        ["Невалидные коды", "Для 10...15 все выходы остаются в '1'."],
        ["GHDL", "TEST PASSED на 160 ns."],
        ["Quartus II", "Full Compilation successful, 0 errors."],
        ["Vivado", "synth_design completed successfully, 0 errors."],
    ], [5.8, 9.8])

    h(doc, "9 Вывод")
    p(doc, "В ходе лабораторной работы разработана VHDL-модель комбинационного устройства варианта 4 - BCD-десятичного дешифратора 145 с активными-низкими выходами.")
    p(doc, "Устройство реализовано двумя способами: через чистые логические выражения и через последовательное описание process/case. Полный перебор входных комбинаций подтвердил корректность обеих моделей. Проект успешно обработан в Quartus II и Xilinx Vivado; в отчетах синтеза регистры отсутствуют, что подтверждает комбинационную природу схемы.")

    doc.add_page_break()
    h(doc, "Приложение А. Листинги VHDL-модулей")
    code(doc, "А.1 Реализация через логические операторы", ROOT / "src" / "sn74145_concurrent.vhd")
    code(doc, "А.2 Реализация через последовательные операторы", ROOT / "src" / "sn74145_sequential.vhd")
    code(doc, "А.3 Верхний уровень сравнения", ROOT / "src" / "sn74145_compare_top.vhd")
    doc.add_page_break()
    h(doc, "Приложение Б. Листинг testbench")
    code(doc, "Б.1 tb_sn74145", ROOT / "tb" / "tb_sn74145.vhd")
    doc.add_page_break()
    h(doc, "Приложение В. Файлы проектов САПР")
    code(doc, "В.1 Quartus QSF", ROOT / "quartus" / "sn74145.qsf")
    code(doc, "В.2 Vivado TCL", ROOT / "vivado" / "create_project.tcl")


def main():
    REPORT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    title_page(doc)
    body(doc)
    out = REPORT / "250541_LR1_SN74145_Власов_РЕ_АПЦУ.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
