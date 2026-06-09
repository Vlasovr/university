from __future__ import annotations

import math
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Результат"
TEMPLATE = ROOT / "Вадима" / "Записка.docx"
DOCX_OUT = OUT / "250541_Власов_РЕ_вариант_4_ОбКС_ПЗ.docx"


def font(size: int, bold: bool = False, italic: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold Italic.ttf" if bold and italic else "",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "",
        "/System/Library/Fonts/Supplemental/Arial Italic.ttf" if italic else "",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for p in candidates:
        if p and Path(p).exists():
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()


FONTS = {
    "xs": font(16),
    "s": font(20),
    "si": font(20, italic=True),
    "m": font(24),
    "mi": font(24, italic=True),
    "b": font(28, True),
    "h": font(34, True),
    "title": font(40, True),
}


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    cur = ""
    for word in words:
        test = word if not cur else f"{cur} {word}"
        if draw.textbbox((0, 0), test, font=fnt)[2] <= width:
            cur = test
        else:
            if cur:
                lines.append(cur)
            cur = word
    if cur:
        lines.append(cur)
    return lines


def centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt, fill="black"):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total_h = sum(heights) + 6 * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + 6


def rect(draw, xy, width=3, dash: bool = False):
    if not dash:
        draw.rectangle(xy, outline="black", width=width)
        return
    x1, y1, x2, y2 = xy
    dashed_line(draw, (x1, y1), (x2, y1), width)
    dashed_line(draw, (x2, y1), (x2, y2), width)
    dashed_line(draw, (x2, y2), (x1, y2), width)
    dashed_line(draw, (x1, y2), (x1, y1), width)


def dashed_line(draw, p1, p2, width=3, dash=18, gap=12):
    x1, y1 = p1
    x2, y2 = p2
    length = math.hypot(x2 - x1, y2 - y1)
    if length == 0:
        return
    dx = (x2 - x1) / length
    dy = (y2 - y1) / length
    pos = 0
    while pos < length:
        end = min(pos + dash, length)
        draw.line((x1 + dx * pos, y1 + dy * pos, x1 + dx * end, y1 + dy * end), fill="black", width=width)
        pos += dash + gap


def dashed_polyline(draw, points, width=3, dash=18, gap=12):
    for p1, p2 in zip(points, points[1:]):
        dashed_line(draw, p1, p2, width=width, dash=dash, gap=gap)


def arrow_line(draw, p1, p2, width=3, dashed=False):
    if dashed:
        dashed_line(draw, p1, p2, width)
    else:
        draw.line((p1[0], p1[1], p2[0], p2[1]), fill="black", width=width)
    ang = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
    for a in (ang + math.pi * 0.82, ang - math.pi * 0.82):
        draw.line((p2[0], p2[1], p2[0] + 18 * math.cos(a), p2[1] + 18 * math.sin(a)), fill="black", width=width)


def title_block(draw: ImageDraw.ImageDraw, title: str, code_suffix: str):
    W, H = 2480, 1754
    rect(draw, (70, 55, W - 70, H - 55), width=5)
    tb = (1320, 1380, W - 90, H - 90)
    rect(draw, tb, width=4)
    x1, y1, x2, y2 = tb
    for y in [1440, 1505, 1570, 1625]:
        draw.line((x1, y, x2, y), fill="black", width=2)
    for x in [1420, 1540, 1760, 1900, 2160, 2275]:
        draw.line((x, y1, x, y2), fill="black", width=2)
    centered(draw, (1760, 1380, 2160, 1505), f"ГУИР.400201.004 {code_suffix}", FONTS["b"])
    centered(draw, (1760, 1505, 2275, 1625), f"Проектирование локальной\nкомпьютерной сети.\n{title}", FONTS["xs"])
    centered(draw, (2275, 1505, x2, 1570), "Лит  Масса  Масштаб", FONTS["xs"])
    centered(draw, (2275, 1625, x2, y2), "Лист 1  Листов 1", FONTS["xs"])
    centered(draw, (1540, 1625, 2275, y2), "ЭВМ, гр. 250541", FONTS["s"])
    labels = ["Изм", "Лист", "№ докум.", "Подпись", "Дата"]
    xx = [1325, 1370, 1425, 1545, 1765]
    for label, x in zip(labels, xx):
        draw.text((x, 1400), label, font=FONTS["xs"], fill="black")
    names = [("Разраб.", "Власов"), ("Пров.", "Смирнов"), ("Т. контр.", ""), ("Реценз.", ""), ("Н. контр.", ""), ("Утверд.", "")]
    y = 1452
    for role, name in names:
        draw.text((1330, y), role, font=FONTS["xs"], fill="black")
        draw.text((1545, y), name, font=FONTS["xs"], fill="black")
        y += 30
    rotated = Image.new("RGBA", (470, 70), (255, 255, 255, 0))
    rd = ImageDraw.Draw(rotated)
    rd.rectangle((0, 0, 470, 70), outline="black", width=4)
    centered(rd, (0, 0, 470, 70), f"ГУИР.400201.004 {code_suffix}", FONTS["h"])
    draw.bitmap((115, 55), rotated.rotate(180, expand=True), fill=None)


def simple_box(draw, xy, text, fnt=None, dash=False):
    fnt = fnt or FONTS["mi"]
    rect(draw, xy, width=3, dash=dash)
    x1, y1, x2, y2 = xy
    lines = []
    for part in text.split("\n"):
        lines.extend(wrap(draw, part, fnt, x2 - x1 - 24))
    centered(draw, xy, "\n".join(lines), fnt)


def draw_router(draw, x, y, label):
    rect(draw, (x, y, x + 160, y + 55), width=3)
    for i in range(5):
        rect(draw, (x + 16 + i * 25, y + 20, x + 32 + i * 25, y + 35), width=2)
    draw.text((x, y - 48), label, font=FONTS["xs"], fill="black")


def draw_switch(draw, x, y, label):
    rect(draw, (x, y, x + 185, y + 55), width=3)
    for row in range(2):
        for i in range(8):
            rect(draw, (x + 15 + i * 20, y + 12 + row * 22, x + 29 + i * 20, y + 25 + row * 22), width=2)
    draw.text((x, y - 44), label, font=FONTS["xs"], fill="black")


def draw_pc(draw, x, y, label):
    rect(draw, (x, y, x + 62, y + 45), width=3)
    draw.line((x + 20, y + 57, x + 42, y + 57), fill="black", width=3)
    draw.line((x + 31, y + 45, x + 31, y + 57), fill="black", width=3)
    rect(draw, (x + 72, y + 8, x + 103, y + 58), width=3)
    draw.text((x - 10, y + 70), label, font=FONTS["xs"], fill="black")


def draw_printer(draw, x, y, label):
    rect(draw, (x, y + 20, x + 80, y + 70), width=3)
    rect(draw, (x + 15, y, x + 65, y + 30), width=2)
    draw.line((x + 20, y + 50, x + 60, y + 50), fill="black", width=2)
    draw.text((x - 15, y + 80), label, font=FONTS["xs"], fill="black")


def draw_ap(draw, x, y, label):
    draw.arc((x, y, x + 90, y + 90), 200, 340, fill="black", width=3)
    draw.arc((x + 15, y + 16, x + 75, y + 76), 205, 335, fill="black", width=3)
    draw.arc((x + 30, y + 32, x + 60, y + 62), 210, 330, fill="black", width=3)
    draw.ellipse((x + 38, y + 70, x + 52, y + 84), fill="black")
    draw.text((x - 5, y + 90), label, font=FONTS["xs"], fill="black")


def draw_ap_plan(draw, x, y, label):
    draw.arc((x, y, x + 66, y + 66), 200, 340, fill="black", width=3)
    draw.arc((x + 12, y + 13, x + 54, y + 55), 205, 335, fill="black", width=3)
    draw.arc((x + 24, y + 26, x + 42, y + 44), 210, 330, fill="black", width=2)
    draw.ellipse((x + 28, y + 51, x + 38, y + 61), fill="black")
    draw.text((x + 5, y + 65), label, font=FONTS["xs"], fill="black")


def draw_server(draw, x, y, label):
    rect(draw, (x, y, x + 70, y + 120), width=3)
    for yy in [y + 28, y + 58, y + 88]:
        draw.line((x + 12, yy, x + 58, yy), fill="black", width=2)
    draw.text((x - 15, y + 132), label, font=FONTS["xs"], fill="black")


def legend_box(draw, items):
    x, y, w, h = 110, 1285, 900, 300
    rect(draw, (x, y, x + w, y + h), width=3)
    draw.text((x + 20, y + 20), "Условные обозначения:", font=FONTS["m"], fill="black")
    col_w = w // 2
    for idx, (kind, text) in enumerate(items):
        col = idx // 3
        row = idx % 3
        xx = x + 25 + col * col_w
        yy = y + 70 + row * 70
        if kind == "pc":
            rect(draw, (xx, yy, xx + 45, yy + 32), width=2)
            rect(draw, (xx + 55, yy + 6, xx + 78, yy + 42), width=2)
        elif kind == "router":
            rect(draw, (xx, yy + 10, xx + 95, yy + 42), width=2)
            for i in range(5):
                rect(draw, (xx + 10 + i * 16, yy + 22, xx + 20 + i * 16, yy + 32), width=1)
        elif kind == "switch":
            rect(draw, (xx, yy + 8, xx + 120, yy + 45), width=2)
            for r in range(2):
                for i in range(6):
                    rect(draw, (xx + 10 + i * 17, yy + 15 + r * 16, xx + 20 + i * 17, yy + 25 + r * 16), width=1)
        elif kind == "ap":
            draw.arc((xx, yy, xx + 70, yy + 70), 205, 335, fill="black", width=3)
            draw.arc((xx + 14, yy + 15, xx + 56, yy + 57), 205, 335, fill="black", width=2)
            draw.ellipse((xx + 31, yy + 55, xx + 41, yy + 65), fill="black")
        elif kind == "printer":
            rect(draw, (xx, yy + 20, xx + 65, yy + 55), width=2)
            rect(draw, (xx + 12, yy + 5, xx + 53, yy + 27), width=2)
        elif kind == "server":
            rect(draw, (xx, yy, xx + 50, yy + 62), width=2)
            for off in [16, 32, 48]:
                draw.line((xx + 9, yy + off, xx + 41, yy + off), fill="black", width=1)
        elif kind == "copper":
            draw.line((xx, yy + 30, xx + 120, yy + 30), fill="black", width=4)
        elif kind == "vpn":
            dashed_line(draw, (xx, yy + 30), (xx + 120, yy + 30), width=4)
        ty = yy + 18
        for line in wrap(draw, text, FONTS["xs"], col_w - 175):
            draw.text((xx + 145, ty), line, font=FONTS["xs"], fill="black")
            ty += 19


def make_canvas():
    img = Image.new("RGB", (2480, 1754), "white")
    return img, ImageDraw.Draw(img)


def save_img(img, path: Path):
    img.save(path, "PNG")


def draw_structural():
    img, draw = make_canvas()
    title_block(draw, "Схема СКС структурная", "С1")
    simple_box(draw, (1080, 150, 1390, 280), "Internet\nGigabit Ethernet", FONTS["mi"])
    rect(draw, (150, 175, 1010, 1160), width=3, dash=True)
    rect(draw, (1470, 175, 2330, 1160), width=3, dash=True)
    draw.text((175, 130), "Первое здание", font=FONTS["b"], fill="black")
    draw.text((1495, 130), "Второе здание", font=FONTS["b"], fill="black")

    simple_box(draw, (560, 245, 940, 500), "Сетевой шкаф 1\nRT1 MikroTik RB5009\nSW1 D-Link DGS-1100-18PV2\nFS1 QNAP TS-233", FONTS["s"], dash=True)
    simple_box(draw, (1580, 245, 1960, 500), "Сетевой шкаф 2\nRT2 MikroTik RB5009\nSW2 D-Link DGS-1100-18PV2", FONTS["s"], dash=True)
    arrow_line(draw, (940, 370), (1080, 215), width=4)
    arrow_line(draw, (1390, 215), (1580, 370), width=4)

    simple_box(draw, (245, 600, 610, 940), "Лаборатория 1\n4 ПК\n10 смартфонов\nPR1 цветной принтер\nAP1", FONTS["s"], dash=True)
    simple_box(draw, (660, 620, 990, 900), "Кабинет директора\n1 ПК\n1 смартфон", FONTS["s"], dash=True)
    simple_box(draw, (1495, 600, 1865, 940), "Лаборатория 2\n6 ПК\n10 смартфонов\nPR2 цветной принтер\nAP2", FONTS["s"], dash=True)
    simple_box(draw, (1915, 620, 2250, 900), "Бухгалтерия\n1 ПК\n1 смартфон", FONTS["s"], dash=True)

    for a, b in [((560, 500), (430, 600)), ((940, 500), (830, 620)), ((1770, 500), (1680, 600)), ((1960, 500), (2080, 620))]:
        arrow_line(draw, a, b, width=4)
    simple_box(draw, (720, 975, 960, 1105), "Служебное\nпомещение\nсерверная зона", FONTS["s"])
    arrow_line(draw, (755, 500), (840, 975), width=4)

    legend_box(draw, [("router", "Маршрутизатор"), ("switch", "Коммутатор"), ("server", "Файловый сервер"), ("copper", "Витая пара Cat.5e/Cat.6"), ("vpn", "IPsec VPN поверх Internet")])
    return img


def draw_functional():
    img, draw = make_canvas()
    title_block(draw, "Схема СКС функциональная", "С2")
    draw.text((145, 105), "Первое здание", font=FONTS["b"], fill="black")
    draw.text((1375, 105), "Второе здание", font=FONTS["b"], fill="black")
    rect(draw, (100, 150, 1130, 1215), width=3, dash=True)
    rect(draw, (1350, 150, 2380, 1215), width=3, dash=True)

    draw_router(draw, 520, 255, "RT1 RB5009UG+S+IN\nWAN 198.51.100.2/30\nVLAN 10/20/30/40")
    draw_switch(draw, 500, 475, "SW1 DGS-1100-18PV2\nmgmt 10.4.40.2/28")
    draw_server(draw, 865, 385, "FS1 QNAP TS-233\n10.4.40.10/28")
    draw_router(draw, 1680, 255, "RT2 RB5009UG+S+IN\nWAN 203.0.113.2/30\nVLAN 110/120/130/140")
    draw_switch(draw, 1660, 475, "SW2 DGS-1100-18PV2\nmgmt 10.4.140.2/29")

    centered(draw, (1135, 275, 1350, 395), "Internet\nIPsec IKEv2", FONTS["mi"])
    arrow_line(draw, (680, 280), (1135, 330), width=4, dashed=True)
    arrow_line(draw, (1350, 330), (1680, 280), width=4, dashed=True)
    draw.line((680, 310, 680, 475), fill="black", width=4)
    draw.line((1840, 310, 1840, 475), fill="black", width=4)
    draw.line((685, 505, 865, 505), fill="black", width=4)

    lab1 = (190, 705, 520, 1055)
    dir1 = (615, 725, 1000, 1020)
    lab2 = (1435, 705, 1815, 1055)
    acc2 = (1910, 725, 2295, 1020)
    for box, text in [
        (lab1, "VLAN 10 Лаборатория\nPC1-PC4: 10.4.10.11-14\nPR1: 10.4.30.11\nVLAN 20 Wi-Fi: 10.4.20.0/27"),
        (dir1, "VLAN 10 Директор\nPC5: 10.4.10.21\nWi-Fi сотрудника через AP1"),
        (lab2, "VLAN 110 Лаборатория\nPC6-PC11: 10.4.110.11-16\nPR2: 10.4.130.11\nVLAN 120 Wi-Fi: 10.4.120.0/27"),
        (acc2, "VLAN 110 Бухгалтерия\nPC12: 10.4.110.21\nWi-Fi сотрудника через AP2"),
    ]:
        simple_box(draw, box, text, FONTS["s"], dash=True)

    draw_ap(draw, 480, 1030, "AP1\n10.4.20.2")
    draw_ap(draw, 1780, 1030, "AP2\n10.4.120.2")
    for p in [(500, 530, 350, 705), (610, 530, 810, 725), (1660, 530, 1620, 705), (1785, 530, 2100, 725)]:
        arrow_line(draw, (p[0], p[1]), (p[2], p[3]), width=4)
    draw.line((555, 530, 525, 1030), fill="black", width=4)
    draw.line((1715, 530, 1825, 1030), fill="black", width=4)
    dashed_line(draw, (525, 1100), (335, 985), width=4)
    dashed_line(draw, (1825, 1100), (1630, 985), width=4)

    legend_box(draw, [("router", "Маршрутизатор L3/VPN/firewall"), ("switch", "Управляемый PoE-коммутатор"), ("server", "Файловый сервер/NAS"), ("copper", "Медная линия СКС"), ("vpn", "Зашифрованный IPsec-туннель")])
    return img


def plan_room(draw, scale, origin, x, y, w, h, name, devices=None):
    ox, oy = origin
    xy = (ox + x * scale, oy + y * scale, ox + (x + w) * scale, oy + (y + h) * scale)
    rect(draw, tuple(map(int, xy)), width=4)
    name_font = FONTS["xs"] if w * h <= 4 else FONTS["s"]
    draw.text((xy[0] + 12, xy[1] + 12), name, font=name_font, fill="black")
    if devices:
        for kind, dx, dy, label in devices:
            px = int(ox + (x + dx) * scale)
            py = int(oy + (y + dy) * scale)
            if kind == "pc":
                draw_pc(draw, px, py, label)
            elif kind == "printer":
                draw_printer(draw, px, py, label)
            elif kind == "ap":
                draw_ap_plan(draw, px, py, label)
            elif kind == "rack":
                draw.rectangle((px, py, px + 70, py + 58), outline="black", width=4)
                for yy in [py + 16, py + 31, py + 46]:
                    draw.line((px + 10, yy, px + 60, yy), fill="black", width=2)
                draw.text((px - 6, py + 64), label, font=FONTS["xs"], fill="black")
            elif kind == "socket":
                rect(draw, (px, py, px + 30, py + 30), width=3)
                draw.line((px + 8, py + 15, px + 22, py + 15), fill="black", width=2)
    return xy


def draw_plan():
    img, draw = make_canvas()
    title_block(draw, "План здания. Схема монтажная", "С3")
    draw.text((115, 105), "Первое здание, 10 x 8 м", font=FONTS["b"], fill="black")
    draw.text((1220, 105), "Второе здание, 12 x 8 м", font=FONTS["b"], fill="black")
    scale = 62
    o1 = (145, 190)
    o2 = (1210, 190)
    rect(draw, (o1[0], o1[1], o1[0] + 10 * scale, o1[1] + 8 * scale), width=5)
    rect(draw, (o2[0], o2[1], o2[0] + 12 * scale, o2[1] + 8 * scale), width=5)

    plan_room(draw, scale, o1, 0, 0, 5, 4, "Лаборатория 20 м2", [
        ("pc", 0.5, 0.7, "PC1"), ("pc", 2.0, 0.7, "PC2"), ("pc", 0.5, 2.4, "PC3"), ("pc", 2.0, 2.4, "PC4"),
        ("printer", 3.65, 2.05, "PR1"), ("ap", 2.45, 1.15, "AP1")
    ])
    plan_room(draw, scale, o1, 5, 0, 4, 2.5, "Кабинет директора 10 м2", [("pc", 1.0, 1.0, "PC5")])
    plan_room(draw, scale, o1, 0, 4, 2.5, 2, "Служ. 5 м2", [])
    plan_room(draw, scale, o1, 5, 4, 4, 2.5, "Служ./серверная 10 м2", [("rack", 2.4, 0.8, "ШК1")])
    draw.text((o1[0] + 210, o1[1] + 6 * scale + 35), "Коридор/проходы и ввод провайдера", font=FONTS["s"], fill="black")

    plan_room(draw, scale, o2, 0, 0, 8, 5, "Лаборатория 40 м2", [
        ("pc", 0.5, 0.7, "PC6"), ("pc", 2.0, 0.7, "PC7"), ("pc", 3.5, 0.7, "PC8"),
        ("pc", 0.5, 3.0, "PC9"), ("pc", 2.0, 3.0, "PC10"), ("pc", 3.5, 3.0, "PC11"),
        ("printer", 6.5, 2.0, "PR2"), ("ap", 4.35, 2.05, "AP2")
    ])
    plan_room(draw, scale, o2, 8, 0, 4, 2.5, "Бухгалтерия 10 м2", [("pc", 1.0, 1.0, "PC12")])
    plan_room(draw, scale, o2, 8, 5, 1.5, 2, "Служ. 3 м2", [("rack", 0.25, 0.65, "ШК2")])
    draw.text((o2[0] + 155, o2[1] + 6 * scale + 64), "Коридор/проходы и ввод провайдера", font=FONTS["s"], fill="black")

    # Cable routes
    for start, end in [
        ((o1[0] + 5 * scale + 3.2 * scale, o1[1] + 4.8 * scale), (o1[0] + 2.7 * scale, o1[1] + 2.0 * scale)),
        ((o1[0] + 5 * scale + 3.2 * scale, o1[1] + 4.8 * scale), (o1[0] + 6.2 * scale, o1[1] + 1.4 * scale)),
    ]:
        dashed_line(draw, (int(start[0]), int(start[1])), (int(end[0]), int(end[1])), width=4, dash=14, gap=10)
    dashed_polyline(draw, [
        (int(o2[0] + 8.65 * scale), int(o2[1] + 6.05 * scale)),
        (int(o2[0] + 7.9 * scale), int(o2[1] + 5.0 * scale)),
        (int(o2[0] + 4.65 * scale), int(o2[1] + 2.25 * scale)),
    ], width=4, dash=14, gap=10)
    dashed_polyline(draw, [
        (int(o2[0] + 9.2 * scale), int(o2[1] + 1.5 * scale)),
        (int(o2[0] + 9.95 * scale), int(o2[1] + 4.95 * scale)),
        (int(o2[0] + 9.95 * scale), int(o2[1] + 6.2 * scale)),
        (int(o2[0] + 8.65 * scale), int(o2[1] + 6.2 * scale)),
    ], width=4, dash=14, gap=10)

    tech = [
        "Технические требования:",
        "1. Горизонтальная подсистема выполняется кабелем U/UTP Cat.5e LSZH; длина постоянной линии не более 90 м.",
        "2. Кабель прокладывается в пластиковом коробе по коридору на высоте 2,3 м; спуски к розеткам выполняются вертикально.",
        "3. Информационные розетки RJ-45 устанавливаются у рабочих мест на высоте 0,3 м, точки доступа - на потолке.",
        "4. Сетевые шкафы ШК1 и ШК2 размещаются в служебных помещениях с ограниченным доступом.",
        "5. Линии питания и СКС прокладываются раздельно; пересечения выполняются под прямым углом.",
        "6. Все кабели маркируются с двух сторон согласно номеру помещения и номеру порта патч-панели.",
    ]
    rect(draw, (1170, 1000, 2380, 1285), width=3)
    yy = 1020
    for line in tech:
        for wline in wrap(draw, line, FONTS["xs"], 1160):
            draw.text((1190, yy), wline, font=FONTS["xs"], fill="black")
            yy += 24
    legend_box(draw, [("pc", "Персональный компьютер"), ("printer", "Цветной сетевой принтер"), ("ap", "Точка доступа Wi-Fi"), ("router", "Сетевой шкаф/активное оборудование"), ("vpn", "Маршрут прокладки кабеля")])
    return img


def set_cell_text(cell, text: str, bold: bool = False, size: int = 11):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(10)


def clear_document(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def p(doc, text="", style=None, align=None, bold=False, italic=False):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    para.paragraph_format.first_line_indent = Cm(1.25) if style in (None, "Normal") and text else None
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    return para


def set_outline_level(para, level: int):
    p_pr = para._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def h1(doc, text):
    para = p(doc, text.upper(), style="Раздел", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    set_outline_level(para, 0)
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para


def h2(doc, text):
    para = p(doc, text, style="Подраздел", bold=True)
    set_outline_level(para, 1)
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(4)
    return para


def center(doc, text, bold=False, size=14):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    r = para.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    return para


def page_break(doc):
    doc.add_page_break()


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    style_table(table)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell_text(table.cell(i, j), str(val), bold=(i == 0), size=10)
    return table


def add_bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Paragraph")
        para.paragraph_format.left_indent = Cm(1.25)
        para.paragraph_format.first_line_indent = Cm(-0.6)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


def add_image_page(doc, img_path: Path, title: str):
    center(doc, title, bold=True, size=14)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(img_path), width=Cm(16.4))


def add_toc_placeholder(doc):
    center(doc, "СОДЕРЖАНИЕ", bold=True)
    toc_entries = [
        (0, "1 ВВЕДЕНИЕ", 4),
        (0, "2 ОБЗОР ЛИТЕРАТУРЫ", 5),
        (1, "2.1 Технология Gigabit Ethernet", 5),
        (1, "2.2 IPv4 и IPv6", 5),
        (1, "2.3 VLAN и логическая сегментация", 6),
        (1, "2.4 IPsec VPN между зданиями", 6),
        (1, "2.5 Беспроводная сеть Wi-Fi", 6),
        (1, "2.6 Файловый сервер", 7),
        (0, "3 СТРУКТУРНОЕ ПРОЕКТИРОВАНИЕ", 8),
        (1, "3.1 Исходная архитектура объекта", 8),
        (1, "3.2 Состав сети", 8),
        (1, "3.3 Структурная схема", 8),
        (0, "4 ФУНКЦИОНАЛЬНОЕ ПРОЕКТИРОВАНИЕ", 10),
        (1, "4.1 Расчёт портовой ёмкости", 10),
        (1, "4.2 Выбор способа связи между зданиями", 10),
        (1, "4.3 Выбор маршрутизатора", 10),
        (1, "4.4 Выбор коммутатора", 11),
        (1, "4.5 Выбор рабочих станций", 11),
        (1, "4.6 Выбор файлового сервера", 11),
        (1, "4.7 Выбор беспроводных точек доступа и принтеров", 12),
        (1, "4.8 Схема адресации", 12),
        (1, "4.9 Политика безопасности и конфигурирование", 13),
        (0, "5 ПРОЕКТИРОВАНИЕ СТРУКТУРИРОВАННОЙ КАБЕЛЬНОЙ СИСТЕМЫ", 14),
        (1, "5.1 План здания и размещение оборудования", 14),
        (1, "5.2 Выбор кабельной среды", 14),
        (1, "5.3 Расчёт кабельной продукции", 14),
        (1, "5.4 Расчёт покрытия Wi-Fi", 15),
        (1, "5.5 Перечень оборудования и материалов", 15),
        (0, "6 ЗАКЛЮЧЕНИЕ", 17),
        (0, "7 СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 18),
        (0, "ПРИЛОЖЕНИЕ А. Схема СКС структурная", 19),
        (0, "ПРИЛОЖЕНИЕ Б. Схема СКС функциональная", 20),
        (0, "ПРИЛОЖЕНИЕ В. План здания. Схема монтажная", 21),
        (0, "ПРИЛОЖЕНИЕ Г. Перечень оборудования, изделий и материалов", 22),
        (0, "ПРИЛОЖЕНИЕ Д. Ведомость документов", 23),
    ]
    for level, title, page_num in toc_entries:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.7 if level else 0)
        para.paragraph_format.first_line_indent = None
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.tab_stops.add_tab_stop(Cm(17.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        for text in (title, "\t", str(page_num)):
            run = para.add_run()
            if text == "\t":
                run.add_tab()
            else:
                run.text = text
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)


def insert_toc_field(docx_path: Path):
    doc = Document(str(docx_path))
    target = None
    for para in doc.paragraphs:
        if para.text.strip() == "[[TOC]]":
            target = para
            break
    if target is None:
        return
    for r in list(target.runs)[::-1]:
        target._p.remove(r._r)
    run = target.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Содержание обновляется в Word"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in [fld_begin, instr, fld_sep, t, fld_end]:
        run._r.append(el)
    doc.save(str(docx_path))

    tmp = OUT / "_toc_patch"
    if tmp.exists():
        shutil.rmtree(tmp)
    tmp.mkdir()
    with zipfile.ZipFile(docx_path, "r") as z:
        z.extractall(tmp)
    settings = tmp / "word" / "settings.xml"
    tree = etree.parse(str(settings))
    root = tree.getroot()
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    uf = root.find(f"{{{ns}}}updateFields")
    if uf is None:
        uf = etree.Element(f"{{{ns}}}updateFields")
        root.insert(0, uf)
    uf.set(f"{{{ns}}}val", "true")
    tree.write(str(settings), xml_declaration=True, encoding="UTF-8", standalone="yes")
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in tmp.rglob("*"):
            if f.is_file():
                z.write(f, f.relative_to(tmp).as_posix())
    shutil.rmtree(tmp)


def build_doc(drawings: dict[str, Path]):
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    for sec in doc.sections:
        sec.top_margin = Cm(1.6)
        sec.bottom_margin = Cm(1.6)
        sec.left_margin = Cm(2.4)
        sec.right_margin = Cm(1.2)

    center(doc, "Министерство образования Республики Беларусь", size=14)
    center(doc, "Учреждение образования", size=14)
    center(doc, "БЕЛОРУССКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ", bold=True, size=14)
    center(doc, "ИНФОРМАТИКИ И РАДИОЭЛЕКТРОНИКИ", bold=True, size=14)
    p(doc)
    p(doc, "Факультет компьютерных систем и сетей")
    p(doc, "Кафедра электронных вычислительных машин")
    p(doc, "Дисциплина: Оборудование компьютерных сетей")
    for _ in range(5):
        p(doc)
    center(doc, "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", bold=True, size=16)
    center(doc, "к курсовому проекту", size=14)
    center(doc, "на тему", size=14)
    p(doc)
    center(doc, "ЛОКАЛЬНАЯ КОМПЬЮТЕРНАЯ СЕТЬ", bold=True, size=16)
    center(doc, "вариант 4", bold=True, size=14)
    p(doc)
    center(doc, "БГУИР КП 1-40 02 01 004 ПЗ", size=14)
    for _ in range(5):
        p(doc)
    p(doc, "Студент гр. 250541    Р.Е. Власов")
    p(doc, "Руководитель          Ю.В. Смирнов")
    for _ in range(5):
        p(doc)
    center(doc, "Минск 2026", size=14)
    page_break(doc)

    center(doc, "ЗАДАНИЕ", bold=True, size=16)
    center(doc, "по курсовому проектированию", size=14)
    p(doc, "Студенту Власову Р.Е.")
    p(doc, "1. Тема проекта: Локальная компьютерная сеть.")
    p(doc, "2. Срок сдачи студентом законченного проекта: 31 декабря 2026 г.")
    p(doc, "3. Исходные данные к проекту:")
    add_bullets(doc, [
        "сфера деятельности - научно-исследовательская организация в области медицины;",
        "объект состоит из двух удаленных друг от друга на 5 км одноэтажных зданий;",
        "первое здание: лаборатория 20 м2 - 4 стационарных и 10 мобильных подключений; кабинет директора 10 м2 - 1 стационарное и 1 мобильное подключение; два служебных помещения 5 и 10 м2;",
        "второе здание: лаборатория 40 м2 - 6 стационарных и 10 мобильных подключений; бухгалтерия 10 м2 - 1 стационарное и 1 мобильное подключение; служебное помещение 3 м2;",
        "оборудование: ПК, личные смартфоны, файловый сервер для обмена файлами и хранения персональных файлов, два цветных принтера;",
        "подключение к Internet осуществляется через Gigabit Ethernet;",
        "адресация: статические IPv4-адреса на внешних интерфейсах и публичные IPv6-подсети от провайдера;",
        "для проектирования принята бюджетная сеть;",
        "безопасность: все сотрудники имеют право выхода в Internet только по протоколу HTTP.",
    ])
    p(doc, "4. Содержание расчетно-пояснительной записки: лист задания, введение, обзор литературы, структурное проектирование, функциональное проектирование, проектирование структурированной кабельной системы, заключение, список использованных источников, приложения.")
    p(doc, "5. Перечень графического материала: схема структурная, схема СКС функциональная, план здания, перечень оборудования, изделий и материалов.")
    p(doc, "Дата выдачи задания: 06.06.2026.")
    page_break(doc)

    add_toc_placeholder(doc)
    page_break(doc)

    h1(doc, "Введение")
    intro = [
        "Локальная компьютерная сеть является базовой инфраструктурой современной научно-исследовательской организации. Для медицинского профиля особенно важны устойчивый доступ к рабочим файлам, централизованное хранение результатов исследований, управляемая печать документов и разделение доверенных рабочих мест от личных мобильных устройств сотрудников. Ошибки на уровне сети приводят не только к снижению производительности, но и к риску нарушения конфиденциальности внутренних материалов.",
        "Целью курсового проекта является разработка проекта локальной компьютерной сети для научно-исследовательской организации в сфере медицины, расположенной в двух одноэтажных зданиях на расстоянии 5 км друг от друга. Проект учитывает бюджетный характер сети, необходимость проводной и беспроводной составляющих, подключение к Internet через Gigabit Ethernet, использование IPv4 и IPv6, а также требование ограничить доступ сотрудников во внешнюю сеть только протоколом HTTP.",
        "В проекте рассматриваются структурная и функциональная архитектура сети, схема адресации, выбор активного и пассивного оборудования, расчёт портовой ёмкости, расчёт кабельной продукции и проверка покрытия беспроводной сети. Оборудование выбрано из моделей, имеющих актуальные предложения в белорусских каталогах либо поставляемых через региональных продавцов, а технические характеристики проверены по данным производителей.",
        "Основной принцип проектирования состоит в разделении задач по уровням: маршрутизация и безопасность выполняются маршрутизаторами, коммутация и питание точек доступа - управляемыми PoE-коммутаторами, хранение файлов - отдельным NAS-сервером, а кабельная система выполняется как самостоятельная структурированная подсистема с запасом портов и маркировкой линий.",
    ]
    for text in intro:
        p(doc, text)

    h1(doc, "1 Обзор литературы")
    sections = {
        "1.1 Технология Gigabit Ethernet": [
            "Gigabit Ethernet является развитием семейства Ethernet и обеспечивает номинальную скорость передачи 1 Гбит/с. Для офисных и исследовательских локальных сетей наиболее распространён вариант 1000BASE-T, работающий по четырём парам медного кабеля категории 5e и выше на расстоянии до 100 м. Он сохраняет совместимость с Ethernet-кадрами и позволяет подключать рабочие станции, серверы и сетевые принтеры без применения специализированных интерфейсов.",
            "В рассматриваемом проекте Gigabit Ethernet применяется как единая скорость проводного доступа: рабочие станции, принтеры, точки доступа и файловый сервер подключаются к коммутаторам по гигабитным портам. Такой подход достаточен для обмена офисными документами, результатов лабораторных измерений и централизованного доступа к файловому серверу, а также не требует дорогой 10-гигабитной инфраструктуры.",
            "Для внешнего подключения к Internet также используется интерфейс Gigabit Ethernet. Это означает, что провайдер предоставляет электрический или оптический Ethernet-ввод, а маршрутизатор организации должен иметь гигабитный WAN-порт, поддержку статической IPv4-адресации, IPv6 и механизмов фильтрации трафика.",
        ],
        "1.2 IPv4 и IPv6": [
            "IPv4 остаётся наиболее распространённым протоколом сетевого уровня в локальных сетях. Он использует 32-битные адреса, поэтому внутри организаций обычно применяются частные диапазоны с трансляцией NAT на внешнем маршрутизаторе. В данном проекте внутренние IPv4-подсети выбраны из частного диапазона 10.0.0.0/8, а статические адреса провайдера используются на WAN-интерфейсах маршрутизаторов.",
            "IPv6 использует 128-битную адресацию и позволяет выделять отдельные подсети /64 для каждого логического сегмента. В проекте IPv6 применяется параллельно с IPv4: провайдер условно выделяет каждому зданию собственный префикс /56, из которого формируются подсети для проводного сегмента, беспроводного сегмента, принтеров и управления.",
            "Совместное применение IPv4 и IPv6 требует одинаковой политики безопасности для обеих версий протокола. Поэтому правила межсетевого экрана и ограничения HTTP-only задаются отдельно для IPv4 и IPv6, чтобы обход блокировок через альтернативный стек был невозможен.",
        ],
        "1.3 VLAN и логическая сегментация": [
            "VLAN позволяет разделять одну физическую коммутационную инфраструктуру на несколько независимых логических сетей. Кадры разных VLAN маркируются по стандарту IEEE 802.1Q на транковых линиях и обрабатываются как отдельные широковещательные домены. Это снижает объём фонового трафика и позволяет назначать разные правила безопасности для групп устройств.",
            "Для медицинской научной организации сегментация важна из-за смешанного состава устройств. Стационарные ПК являются доверенными рабочими местами, а личные смартфоны сотрудников нельзя считать равноправными участниками внутренней сети. Поэтому мобильные устройства выделены в отдельные Wi-Fi VLAN с доступом только к DNS и HTTP во внешнюю сеть, без доступа к файловому серверу и принтерам.",
            "Принтеры и сетевое управление также вынесены в отдельные VLAN. Это облегчает контроль доступа: рабочие станции могут печатать на принтерах, но мобильные устройства и внешние узлы не получают прямого доступа к служебным интерфейсам оборудования.",
        ],
        "1.4 IPsec VPN между зданиями": [
            "Здания организации удалены друг от друга на 5 км. Прокладка собственного волоконно-оптического кабеля для бюджетной сети экономически нецелесообразна, а радиомост требует прямой видимости, согласования размещения антенн и регулярного обслуживания наружного оборудования. Поэтому связь между зданиями проектируется как маршрутизируемый IPsec VPN поверх двух подключений к Internet.",
            "IPsec обеспечивает шифрование и контроль целостности трафика между площадками. Для проекта выбран режим site-to-site с IKEv2, где оба маршрутизатора имеют статические IPv4-адреса провайдера. Через туннель передаются только внутренние подсети организации: доступ к файловому серверу, печати и служебному управлению.",
            "Преимущество такого решения состоит в низких капитальных затратах и независимости от физической трассы между зданиями. Недостатком является зависимость от доступности Internet-каналов провайдера, поэтому в записке дополнительно предусматривается журналирование и возможность дальнейшего подключения резервного канала.",
        ],
        "1.5 Беспроводная сеть Wi-Fi": [
            "Беспроводной сегмент предназначен для личных смартфонов сотрудников. В проекте он не заменяет проводное подключение рабочих мест, а дополняет его: все основные стационарные компьютеры, принтеры и файловый сервер подключаются по кабелю, что обеспечивает предсказуемую производительность и упрощает контроль безопасности.",
            "Для покрытия помещений выбраны потолочные точки доступа Wi-Fi 5 класса AC1200. В каждом здании устанавливается одна точка доступа, поскольку площади помещений невелики, а расчёт RSSI показывает запас по уровню сигнала даже в диапазоне 5 ГГц. Управление SSID и VLAN выполняется на точках доступа и маршрутизаторах.",
            "Беспроводная сеть разделена от проводной сети сотрудников. Это особенно важно для медицинской организации, где на личных устройствах нельзя гарантировать уровень защиты, наличие антивируса и соблюдение политики хранения данных.",
        ],
        "1.6 Файловый сервер": [
            "Файловый сервер необходим для обмена рабочими документами и хранения персональных файлов сотрудников. Для малой организации с 12 стационарными рабочими местами полноценный стоечный сервер является избыточным по цене, энергопотреблению и обслуживанию. Поэтому в проекте используется двухдисковый NAS с поддержкой SMB и RAID 1.",
            "RAID 1 не заменяет резервное копирование, но защищает от отказа одного диска и позволяет продолжить работу до замены накопителя. Для дальнейшего развития можно подключить внешний диск или облачную репликацию, однако в рамках бюджетного курсового проекта достаточно предусмотреть зеркальный массив и разграничение прав доступа на уровне общих папок.",
        ],
    }
    for title, paras in sections.items():
        h2(doc, title)
        for text in paras:
            p(doc, text)

    h1(doc, "2 Структурное проектирование")
    h2(doc, "2.1 Исходная архитектура объекта")
    for text in [
        "Объект состоит из двух одноэтажных зданий, удалённых друг от друга на 5 км. В первом здании расположены лаборатория, кабинет директора и два служебных помещения. Во втором здании расположены лаборатория, бухгалтерия и одно служебное помещение. Для планирования СКС приняты прямоугольные планы зданий: первое здание 10 x 8 м, второе здание 12 x 8 м.",
        "В первом здании служебное помещение площадью 10 м2 используется как серверная зона и место установки сетевого шкафа. В нём размещаются маршрутизатор RT1, коммутатор SW1, патч-панель и файловый сервер FS1. Во втором здании сетевой шкаф ШК2 размещается в служебном помещении 3 м2; в нём находятся маршрутизатор RT2, коммутатор SW2 и патч-панель.",
        "Такое размещение уменьшает длину горизонтальных линий, обеспечивает ограниченный физический доступ к активному оборудованию и позволяет вводить линию провайдера непосредственно в служебное помещение каждого здания.",
    ]:
        p(doc, text)

    h2(doc, "2.2 Состав сети")
    add_bullets(doc, [
        "два маршрутизатора MikroTik RB5009UG+S+IN для подключения к провайдеру, маршрутизации VLAN, IPsec VPN и межсетевого экрана;",
        "два управляемых PoE-коммутатора D-Link DGS-1100-18PV2 для подключения рабочих мест, принтеров, точек доступа и серверного оборудования;",
        "две потолочные точки доступа MikroTik cAP XL ac для беспроводного сегмента личных смартфонов;",
        "двенадцать стационарных ПК: пять в первом здании и семь во втором;",
        "один файловый сервер QNAP TS-233 с двумя дисками WD Red Plus 2 TB в RAID 1;",
        "два цветных сетевых принтера Pantum CP2100DW, по одному в каждой лаборатории;",
        "пассивная СКС на базе медного кабеля U/UTP Cat.5e LSZH, патч-панелей, розеток RJ-45 и кабель-каналов.",
    ])
    h2(doc, "2.3 Структурная схема")
    p(doc, "Структурная схема сети приведена в приложении А. На ней показаны два здания, сетевые шкафы, маршрутизаторы, коммутаторы, файловый сервер, точки доступа, принтеры и группы рабочих мест. Между маршрутизаторами организован IPsec VPN поверх Internet.")
    add_table(doc, [
        ["Блок", "Назначение", "Основное оборудование"],
        ["Маршрутизация", "WAN, NAT, IPv4/IPv6, IPsec, firewall", "2 x MikroTik RB5009UG+S+IN"],
        ["Коммутация", "Access/trunk VLAN, PoE для AP", "2 x D-Link DGS-1100-18PV2"],
        ["Проводной доступ", "Подключение ПК, принтеров и NAS", "U/UTP Cat.5e, RJ-45"],
        ["Беспроводной доступ", "Личные смартфоны сотрудников", "2 x MikroTik cAP XL ac"],
        ["Хранение файлов", "Общие и персональные папки", "QNAP TS-233, RAID 1"],
    ])

    h1(doc, "3 Функциональное проектирование")
    h2(doc, "3.1 Расчёт портовой ёмкости")
    p(doc, "Портовая ёмкость определяется по количеству стационарных проводных подключений и служебных соединений внутри шкафов. Мобильные устройства подключаются через точки доступа и не требуют индивидуальных портов коммутатора.")
    add_table(doc, [
        ["Категория", "Первое здание", "Второе здание", "Итого"],
        ["Рабочие станции", "5", "7", "12"],
        ["Цветные сетевые принтеры", "1", "1", "2"],
        ["Точки доступа Wi-Fi", "1", "1", "2"],
        ["Файловый сервер", "1", "-", "1"],
        ["Соединение маршрутизатор-коммутатор", "1", "1", "2"],
        ["Итого задействовано", "9", "10", "19"],
        ["Резерв портов коммутатора", "7", "6", "13"],
    ])
    p(doc, "Выбранный коммутатор D-Link DGS-1100-18PV2 имеет 16 гигабитных медных портов PoE и 2 SFP/Combo-порта. Даже с учётом служебного соединения с маршрутизатором в каждом здании остаётся не менее шести свободных портов, что достаточно для дальнейшего подключения дополнительных рабочих мест или оборудования.")

    h2(doc, "3.2 Выбор способа связи между зданиями")
    add_table(doc, [
        ["Вариант", "Преимущества", "Недостатки", "Итог"],
        ["Собственная ВОЛС 5 км", "Высокая скорость, независимость от Internet", "Высокая стоимость прокладки, согласования трассы", "Не выбран для бюджетной сети"],
        ["Радиомост", "Нет аренды канала, можно получить высокую скорость при прямой видимости", "Нужны мачты/антенны, влияние погоды, обслуживание наружных узлов", "Не выбран"],
        ["IPsec VPN поверх Internet", "Минимальные капитальные затраты, использует статические адреса провайдера, быстро масштабируется", "Зависит от доступности провайдера", "Выбран"],
    ])
    p(doc, "Для связи между зданиями принят IPsec VPN site-to-site. Для курсового проекта это наиболее рациональное решение: оно соответствует бюджетной сети, не требует прокладки физической трассы между зданиями и обеспечивает шифрование служебного трафика.")

    h2(doc, "3.3 Выбор маршрутизатора")
    add_table(doc, [
        ["Параметр", "MikroTik RB5009UG+S+IN", "MikroTik hEX S RB760iGS"],
        ["WAN/LAN интерфейсы", "7 x 1G, 1 x 2.5G, 1 x SFP+", "5 x 1G, 1 x SFP"],
        ["Производительность routing с фильтрами", "до 9365,6 Мбит/с при 1518 байт", "до 766,4 Мбит/с при 1518 байт"],
        ["IPsec single tunnel", "до 1354,1 Мбит/с при 1400 байт", "до 231,8 Мбит/с при 1400 байт"],
        ["IPv6/VLAN/firewall", "RouterOS v7, поддерживается", "RouterOS, поддерживается"],
        ["Цена в РБ, BYN", "от 799,00", "от 358,25"],
        ["Оценка", "Выбран: обеспечивает запас под Gigabit и VPN", "Дешевле, но слабее для межзданного VPN"],
    ])
    p(doc, "Несмотря на более высокую стоимость, выбран MikroTik RB5009UG+S+IN. Для варианта 4 ключевыми являются Gigabit Ethernet на внешнем канале и межзданный IPsec VPN; младший hEX S подходит для малых офисов, но имеет заметно меньший запас по IPsec.")

    h2(doc, "3.4 Выбор коммутатора")
    add_table(doc, [
        ["Параметр", "D-Link DGS-1100-18PV2", "D-Link DGS-1100-16V2"],
        ["Порты", "16 x 1G PoE + 2 Combo/SFP", "16 x 1G"],
        ["VLAN", "802.1Q, статические VLAN", "802.1Q, статические VLAN"],
        ["PoE", "Есть, бюджет 130 Вт", "Нет"],
        ["Назначение", "Питание AP и коммутация рабочих мест", "Требуются отдельные PoE-инжекторы"],
        ["Цена в РБ, BYN", "от 696,14", "от 326,69"],
        ["Оценка", "Выбран", "Дешевле, но хуже для эксплуатации"],
    ])
    p(doc, "DGS-1100-18PV2 выбран из-за сочетания VLAN, гигабитных портов и PoE. Использование PoE-коммутатора упрощает питание потолочных точек доступа и повышает аккуратность монтажа в небольших помещениях.")

    h2(doc, "3.5 Выбор рабочих станций")
    add_table(doc, [
        ["Параметр", "N-Tech Comfort WH10005", "Jet Office 3i8100D8SD12"],
        ["Процессор", "AMD Ryzen 5 5600G", "Intel Core i3-8100"],
        ["ОЗУ", "16 ГБ", "8 ГБ"],
        ["Накопитель", "SSD 512 ГБ", "SSD 120 ГБ"],
        ["ОС", "Windows 11 Pro", "Без ОС"],
        ["Цена в РБ, BYN", "от 1500,00", "от 955,00"],
        ["Оценка", "Выбран: лучше подходит для рабочих мест НИО", "Дешевле, но требует ОС и модернизации"],
    ])
    p(doc, "Для стационарных рабочих мест выбран N-Tech Comfort WH10005. Он дороже минимального офисного ПК, но имеет 16 ГБ ОЗУ, SSD 512 ГБ и Windows 11 Pro, что снижает скрытые затраты на дооснащение и установку операционной системы.")

    h2(doc, "3.6 Выбор файлового сервера")
    add_table(doc, [
        ["Параметр", "QNAP TS-233", "Полноценный стоечный сервер"],
        ["Форм-фактор", "2-дисковый NAS", "1U/4U сервер"],
        ["Диски", "2 x SATA, RAID 1", "Зависит от конфигурации"],
        ["Сетевой порт", "1 x 1GbE", "1GbE/10GbE"],
        ["Обслуживание", "Простое веб-администрирование", "Требует ОС и администрирования сервера"],
        ["Оценка", "Выбран для бюджетной сети на 12 ПК", "Избыточен по цене"],
    ])
    p(doc, "QNAP TS-233 с двумя дисками WD Red Plus 2 TB в RAID 1 обеспечивает централизованное хранение файлов. Полезный объём после зеркалирования составляет около 2 TB, что достаточно для документов малой исследовательской организации. Для защиты от удаления файлов дополнительно предусматриваются права доступа и периодическое резервное копирование на внешний носитель.")

    h2(doc, "3.7 Выбор беспроводных точек доступа и принтеров")
    add_table(doc, [
        ["Оборудование", "Выбранная модель", "Обоснование"],
        ["Точка доступа", "MikroTik cAP XL ac", "Wi-Fi 5, 2,4/5 ГГц, AC1200, потолочный монтаж, питание PoE, актуальные предложения в РБ"],
        ["Цветной принтер", "Pantum CP2100DW", "Лазерная цветная печать A4, Ethernet и Wi-Fi, 20 стр./мин, автоматическая двусторонняя печать"],
    ])
    p(doc, "Точки доступа используются только для личных смартфонов сотрудников, поэтому Wi-Fi 5 достаточно. Принтеры подключаются по Ethernet в отдельные VLAN принтеров; встроенный Wi-Fi у принтеров в проекте отключается, чтобы избежать обхода политики VLAN.")

    h2(doc, "3.8 Схема адресации")
    p(doc, "Внутри организации применяются частные IPv4-подсети с маршрутизацией между VLAN на маршрутизаторах. Для внешних интерфейсов условно приняты статические адреса из документальных диапазонов RFC 5737; в реальном проекте они заменяются адресами, выданными провайдером. IPv6-подсети также приведены в документальном префиксе 2001:db8::/32 как учебная замена публичных префиксов провайдера.")
    add_table(doc, [
        ["Сегмент", "VLAN", "IPv4", "Шлюз", "IPv6 /64"],
        ["Здание 1: проводные ПК", "10", "10.4.10.0/27", "10.4.10.1", "2001:db8:2505:4a:10::/64"],
        ["Здание 1: Wi-Fi", "20", "10.4.20.0/27", "10.4.20.1", "2001:db8:2505:4a:20::/64"],
        ["Здание 1: принтеры", "30", "10.4.30.0/29", "10.4.30.1", "2001:db8:2505:4a:30::/64"],
        ["Здание 1: сервер/management", "40", "10.4.40.0/28", "10.4.40.1", "2001:db8:2505:4a:40::/64"],
        ["Здание 2: проводные ПК", "110", "10.4.110.0/27", "10.4.110.1", "2001:db8:2505:4b:110::/64"],
        ["Здание 2: Wi-Fi", "120", "10.4.120.0/27", "10.4.120.1", "2001:db8:2505:4b:120::/64"],
        ["Здание 2: принтеры", "130", "10.4.130.0/29", "10.4.130.1", "2001:db8:2505:4b:130::/64"],
        ["Здание 2: management", "140", "10.4.140.0/29", "10.4.140.1", "2001:db8:2505:4b:140::/64"],
    ])
    add_table(doc, [
        ["Устройство", "IPv4", "IPv6", "Примечание"],
        ["RT1 WAN", "198.51.100.2/30", "2001:db8:ffff:1::2/64", "шлюз провайдера 198.51.100.1"],
        ["RT2 WAN", "203.0.113.2/30", "2001:db8:ffff:2::2/64", "шлюз провайдера 203.0.113.1"],
        ["SW1", "10.4.40.2", "2001:db8:2505:4a:40::2", "управление"],
        ["SW2", "10.4.140.2", "2001:db8:2505:4b:140::2", "управление"],
        ["FS1", "10.4.40.10", "2001:db8:2505:4a:40::10", "файловый сервер"],
        ["AP1", "10.4.20.2", "2001:db8:2505:4a:20::2", "точка доступа"],
        ["AP2", "10.4.120.2", "2001:db8:2505:4b:120::2", "точка доступа"],
        ["PR1", "10.4.30.11", "2001:db8:2505:4a:30::11", "принтер"],
        ["PR2", "10.4.130.11", "2001:db8:2505:4b:130::11", "принтер"],
    ])

    h2(doc, "3.9 Политика безопасности и конфигурирование")
    p(doc, "Основная политика безопасности реализуется на маршрутизаторах MikroTik. Коммутаторы выполняют разделение access/trunk-портов и передачу VLAN до маршрутизатора, а маршрутизаторы задают правила маршрутизации, NAT, IPsec и фильтрации.")
    add_bullets(doc, [
        "доступ из WAN во внутренние подсети запрещён, кроме служебных пакетов IPsec между RT1 и RT2;",
        "доступ рабочих ПК к Internet разрешён только по TCP/80; DNS разрешён только к локальному маршрутизатору;",
        "доступ личных смартфонов к внутренним VLAN запрещён, кроме DNS к маршрутизатору;",
        "доступ рабочих ПК к FS1 разрешён по SMB TCP/445 через VPN между зданиями;",
        "печать разрешена только из проводных рабочих VLAN к VLAN принтеров;",
        "управление маршрутизаторами и коммутаторами доступно только из management VLAN и запрещено из Wi-Fi.",
    ])
    p(doc, "Пример логики правил RouterOS: accept established,related; accept IPsec peer UDP/500, UDP/4500 и ESP; accept DNS от внутренних VLAN к адресу маршрутизатора; accept LAN-to-WAN tcp/80; accept wired VLAN к FS1 tcp/445; accept wired VLAN к принтерам; drop прочий inter-VLAN; drop прочий LAN-to-WAN; drop WAN-to-LAN. Для IPv6 создаётся аналогичный набор правил в /ipv6 firewall filter.")

    h1(doc, "4 Проектирование структурированной кабельной системы")
    h2(doc, "4.1 План здания и размещение оборудования")
    p(doc, "План здания и схема монтажной прокладки приведены в приложении В. Кабельные трассы прокладываются в кабель-каналах по коридорам и стенам помещений. Сетевые шкафы установлены в служебных помещениях, что уменьшает риск физического доступа посторонних лиц.")
    p(doc, "В первом здании точка доступа размещается в лаборатории ближе к центру площади, что одновременно покрывает лабораторию, кабинет директора и проходы. Во втором здании точка доступа размещается в лаборатории, так как именно там находится основная группа мобильных устройств; бухгалтерия покрывается за счёт небольшой площади и близости к лаборатории.")

    h2(doc, "4.2 Выбор кабельной среды")
    p(doc, "Для горизонтальной подсистемы выбран медный кабель U/UTP Cat.5e LSZH. Категория 5e поддерживает 1000BASE-T на длине канала до 100 м, а оболочка LSZH уменьшает дымовыделение и выделение галогенов при нагреве, что важно для офисных и медицинских помещений. Использование категории 6 также возможно, но для бюджетного проекта и заданной скорости 1 Гбит/с категория 5e достаточна.")
    p(doc, "Между зданиями медная или волоконная линия СКС не прокладывается: межзданный обмен выполняется через IPsec VPN поверх Internet. Поэтому внутренняя СКС каждого здания является самостоятельной горизонтальной подсистемой.")

    h2(doc, "4.3 Расчёт кабельной продукции")
    p(doc, "В расчёт включены постоянные линии от сетевого шкафа до рабочих мест, принтеров и точек доступа. Сервер и соединения маршрутизатор-коммутатор находятся в шкафу и подключаются короткими патч-кордами, поэтому в расчёт горизонтального кабеля не входят.")
    add_table(doc, [
        ["Здание", "Количество линий", "Средняя горизонтальная длина, м", "Вертикаль и запас, м", "Расчётная длина, м"],
        ["Первое", "7", "7,5", "5,9", "93,8"],
        ["Второе", "9", "9,5", "5,9", "138,6"],
        ["Итого без общего запаса", "16", "-", "-", "232,4"],
        ["Общий монтажный запас 10 %", "-", "-", "-", "23,2"],
        ["Итого к закупке", "-", "-", "-", "255,6"],
    ])
    p(doc, "Стандартная бухта кабеля составляет 305 м. Расчётная потребность 255,6 м укладывается в одну бухту, остаток около 49 м используется как эксплуатационный запас для перезаделки и возможного добавления линий.")

    h2(doc, "4.4 Расчёт покрытия Wi-Fi")
    p(doc, "Проверка выполняется для диапазона 5 ГГц, поскольку он имеет большее затухание по сравнению с 2,4 ГГц. Если расчётный уровень сигнала в 5 ГГц выше -65 дБм, покрытие для офисных задач считается достаточным.")
    p(doc, "Формула потерь в свободном пространстве: Lfs = 32,44 + 20 lg(fМГц) + 20 lg(dкм). Для второго здания принимается максимальная дистанция до клиента 10 м, частота 5180 МГц. Тогда Lfs = 32,44 + 20 lg(5180) + 20 lg(0,01) = 66,7 дБ. С учётом двух внутренних перегородок 15 дБ и запаса 5 дБ полные потери составляют 86,7 дБ.")
    p(doc, "При ограниченной мощности передатчика 20 дБм и усилении антенны 5,5 дБи расчётный RSSI равен 20 + 5,5 - 86,7 = -61,2 дБм. Это выше порогового значения -65 дБм, следовательно одна точка доступа MikroTik cAP XL ac в каждом здании обеспечивает достаточное покрытие для смартфонов сотрудников.")

    h2(doc, "4.5 Перечень оборудования и материалов")
    equipment = [
        ["1", "Маршрутизатор MikroTik RB5009UG+S+IN", "шт.", "2", "799,00", "1598,00"],
        ["2", "Коммутатор D-Link DGS-1100-18PV2", "шт.", "2", "696,14", "1392,28"],
        ["3", "Точка доступа MikroTik cAP XL ac", "шт.", "2", "344,46", "688,92"],
        ["4", "ПК N-Tech Comfort WH10005", "шт.", "12", "1500,00", "18000,00"],
        ["5", "Файловый сервер QNAP TS-233", "шт.", "1", "964,00", "964,00"],
        ["6", "HDD WD Red Plus 2TB WD20EFPX", "шт.", "2", "800,04", "1600,08"],
        ["7", "Принтер Pantum CP2100DW", "шт.", "2", "1547,11", "3094,22"],
        ["8", "Шкаф телекоммуникационный настенный 19\", 9U", "шт.", "2", "350,00", "700,00"],
        ["9", "Патч-панель 24 порта Cat.5e", "шт.", "2", "120,00", "240,00"],
        ["10", "Кабель U/UTP Cat.5e LSZH, бухта 305 м", "бухта", "1", "350,00", "350,00"],
        ["11", "Розетка RJ-45 Cat.5e, внешняя/вставка", "шт.", "16", "15,00", "240,00"],
        ["12", "Патч-корд RJ-45 Cat.5e 1-2 м", "шт.", "40", "6,00", "240,00"],
        ["13", "Кабель-канал и монтажная фурнитура", "компл.", "1", "360,00", "360,00"],
        ["14", "ИБП для сетевого шкафа", "шт.", "2", "330,00", "660,00"],
    ]
    add_table(doc, [["№", "Наименование", "Ед.", "Кол.", "Цена, BYN", "Сумма, BYN"]] + equipment + [["", "Итого ориентировочно", "", "", "", "30127,50"]])
    p(doc, "Цены приняты ориентировочно по состоянию на 08.06.2026 для подтверждения реализуемости закупки в Республике Беларусь. Для пассивных материалов допускается подбор эквивалентных изделий того же класса при сохранении категории кабельной системы и LSZH-оболочки.")

    h1(doc, "Заключение")
    for text in [
        "В ходе курсового проекта разработана локальная компьютерная сеть для научно-исследовательской организации медицинского профиля по варианту 4. Сеть охватывает два одноэтажных здания, удалённых на 5 км, и включает проводную инфраструктуру для рабочих станций, файлового сервера и принтеров, а также беспроводный сегмент для личных смартфонов сотрудников.",
        "Для межзданной связи выбран IPsec VPN поверх двух подключений Gigabit Ethernet к провайдеру. Такое решение соответствует бюджетному характеру проекта и не требует строительства собственной трассы между зданиями. В качестве маршрутизаторов выбраны MikroTik RB5009UG+S+IN, поскольку они имеют достаточный запас производительности для гигабитного канала, фильтрации трафика и IPsec.",
        "Логическая структура сети основана на VLAN. Проводные рабочие места, Wi-Fi, принтеры и управление разделены по разным подсетям. Политика безопасности ограничивает выход сотрудников в Internet протоколом HTTP, запрещает доступ личных смартфонов к внутренним ресурсам и разрешает рабочим ПК доступ к файловому серверу и принтерам только по необходимым протоколам.",
        "Проект СКС предусматривает размещение сетевых шкафов в служебных помещениях, прокладку кабеля U/UTP Cat.5e LSZH, маркировку линий и запас портовой ёмкости. Расчёты показали, что одной бухты 305 м достаточно для горизонтальной подсистемы обоих зданий, а одна точка доступа в каждом здании обеспечивает требуемое покрытие в диапазоне 5 ГГц.",
    ]:
        p(doc, text)

    h1(doc, "Список использованных источников")
    sources = [
        "Таненбаум, Э. С., Фимстер, Н., Уэзеролл, Д. Компьютерные сети. 6-е изд. - Санкт-Петербург: Питер, 2023. - 992 с.",
        "Чаклова, М. И. Проектирование сетей связи. - Минск: БГУИР, 2012. - 95 с.",
        "IEEE Std 802.3. Ethernet Working Group. [Электронный ресурс]. - Режим доступа: https://standards.ieee.org/standard/802_3-2022.html. - Дата доступа: 08.06.2026.",
        "ISO/IEC 11801. Information technology - Generic cabling for customer premises. [Электронный ресурс]. - Режим доступа: https://www.iso.org/standard/66182.html. - Дата доступа: 08.06.2026.",
        "ГОСТ Р 53246-2025. Информационные технологии. Системы кабельные структурированные. Общие технические требования. [Электронный ресурс]. - Режим доступа: https://protect.gost.ru/document.aspx?control=7&id=270937. - Дата доступа: 08.06.2026.",
        "MikroTik RB5009UG+S+IN. Технические характеристики. [Электронный ресурс]. - Режим доступа: https://mikrotik.com/product/rb5009ug_s_in. - Дата доступа: 08.06.2026.",
        "D-Link DGS-1100-18PV2. Технические характеристики. [Электронный ресурс]. - Режим доступа: https://www.dlink.com/en/products/dgs-1100-18pv2-18-port-gigabit-poe-easy-smart-switch. - Дата доступа: 08.06.2026.",
        "MikroTik cAP XL ac. Технические характеристики. [Электронный ресурс]. - Режим доступа: https://mikrotik.com/product/cap_xl_ac. - Дата доступа: 08.06.2026.",
        "QNAP TS-233. Hardware Specs. [Электронный ресурс]. - Режим доступа: https://www.qnap.com/en/product/ts-233/specs. - Дата доступа: 08.06.2026.",
        "Pantum CP2100DW. Технические характеристики. [Электронный ресурс]. - Режим доступа: https://www.pantum.ru/products/laser-devices/cvetnoj-printer-cp2100dw/. - Дата доступа: 08.06.2026.",
        "Каталог Onliner. Сетевое оборудование, ПК, NAS и принтеры. [Электронный ресурс]. - Режим доступа: https://catalog.onliner.by/. - Дата доступа: 08.06.2026.",
    ]
    for i, source in enumerate(sources, 1):
        p(doc, f"{i}. {source}")

    page_break(doc)
    center(doc, "ПРИЛОЖЕНИЕ А", bold=True)
    center(doc, "(обязательное)")
    center(doc, "Схема СКС структурная", bold=True)
    add_image_page(doc, drawings["struct"], "Рисунок А.1 - Схема СКС структурная")
    page_break(doc)
    center(doc, "ПРИЛОЖЕНИЕ Б", bold=True)
    center(doc, "(обязательное)")
    center(doc, "Схема СКС функциональная", bold=True)
    add_image_page(doc, drawings["func"], "Рисунок Б.1 - Схема СКС функциональная")
    page_break(doc)
    center(doc, "ПРИЛОЖЕНИЕ В", bold=True)
    center(doc, "(обязательное)")
    center(doc, "План здания. Схема монтажная", bold=True)
    add_image_page(doc, drawings["plan"], "Рисунок В.1 - План зданий и схема монтажной прокладки")
    page_break(doc)
    center(doc, "ПРИЛОЖЕНИЕ Г", bold=True)
    center(doc, "(обязательное)")
    center(doc, "Перечень оборудования, изделий и материалов", bold=True)
    add_table(doc, [["№", "Наименование", "Ед.", "Кол.", "Примечание"]] + [[r[0], r[1], r[2], r[3], "Закупка/эквивалент"] for r in equipment])
    page_break(doc)
    center(doc, "ПРИЛОЖЕНИЕ Д", bold=True)
    center(doc, "(обязательное)")
    center(doc, "Ведомость документов", bold=True)
    add_table(doc, [
        ["Обозначение", "Наименование", "Листов"],
        ["ГУИР.400201.004 ПЗ", "Пояснительная записка", "1"],
        ["ГУИР.400201.004 С1", "Схема СКС структурная", "1"],
        ["ГУИР.400201.004 С2", "Схема СКС функциональная", "1"],
        ["ГУИР.400201.004 С3", "План здания. Схема монтажная", "1"],
        ["ГУИР.400201.004 ПЭ", "Перечень оборудования, изделий и материалов", "1"],
    ])

    doc.save(str(DOCX_OUT))


def main():
    OUT.mkdir(exist_ok=True)
    drawings = {
        "struct": OUT / "ГУИР.400201.004_С1_Схема_структурная.png",
        "func": OUT / "ГУИР.400201.004_С2_Схема_функциональная.png",
        "plan": OUT / "ГУИР.400201.004_С3_План_здания.png",
    }
    save_img(draw_structural(), drawings["struct"])
    save_img(draw_functional(), drawings["func"])
    save_img(draw_plan(), drawings["plan"])
    build_doc(drawings)
    print(DOCX_OUT)
    for path in drawings.values():
        print(path)


if __name__ == "__main__":
    main()
